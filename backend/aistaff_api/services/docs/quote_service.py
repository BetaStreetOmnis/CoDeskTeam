from __future__ import annotations

from datetime import datetime
import os
from pathlib import Path
import re
import shutil
import subprocess
import tempfile
from uuid import uuid4

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor as DocxRGBColor
from pptx import Presentation
from pptx.dml.color import RGBColor
from pptx.enum.shapes import MSO_SHAPE, MSO_SHAPE_TYPE
from pptx.enum.text import MSO_ANCHOR, PP_ALIGN
from pptx.util import Inches, Pt
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from ...config import Settings
from ...output_cleanup import maybe_cleanup_outputs_dir
from ...url_utils import abs_url
from ..auth_service import create_download_token


class QuoteDocService:
    def __init__(self, settings: Settings) -> None:
        self._settings = settings
        self._settings.outputs_dir.mkdir(parents=True, exist_ok=True)
        maybe_cleanup_outputs_dir(
            self._settings.outputs_dir,
            ttl_seconds=max(0, int(self._settings.outputs_ttl_hours)) * 3600,
        )

    def _apply_doc_theme(self, doc: Document) -> None:
        font_name = "微软雅黑"
        base = doc.styles["Normal"]
        base.font.name = font_name
        base.font.size = Pt(11)
        base._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

        for style_name in ("Heading 1", "Heading 2", "Heading 3"):
            if style_name in doc.styles:
                style = doc.styles[style_name]
                style.font.name = font_name
                style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _set_cell_fill(self, cell, fill: str) -> None:  # noqa: ANN001
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _set_cell_bold(self, cell, *, color: str | None = None) -> None:  # noqa: ANN001
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                if color:
                    run.font.color.rgb = DocxRGBColor.from_string(color)
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(11)

    async def create_quote_docx(
        self,
        *,
        seller: str,
        buyer: str,
        currency: str,
        items: list[dict],
        note: str | None,
    ) -> dict:
        def _fmt_qty(value: float) -> str:
            try:
                v = float(value)
            except Exception:
                return "0"
            if abs(v - int(v)) < 1e-9:
                return str(int(v))
            return f"{v:g}"

        def _fmt_money(value: float) -> str:
            try:
                v = float(value)
            except Exception:
                v = 0.0
            return f"{v:,.2f}"

        doc = Document()
        self._apply_doc_theme(doc)

        # Page setup (slightly tighter margins for business docs)
        section = doc.sections[0]
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)

        quote_no = f"Q-{uuid4().hex[:8].upper()}"
        quote_date = datetime.now().strftime("%Y-%m-%d")

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("报价单")
        title_run.bold = True
        title_run.font.size = Pt(22)
        title_run.font.name = "微软雅黑"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        meta_table = doc.add_table(rows=2, cols=4)
        meta_table.style = "Table Grid"
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_table.autofit = False

        meta_col_widths = [Inches(1.05), Inches(2.65), Inches(1.05), Inches(2.65)]
        for row in meta_table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = meta_col_widths[idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        meta_table.cell(0, 0).text = "供方"
        meta_table.cell(0, 1).text = seller
        meta_table.cell(0, 2).text = "需方"
        meta_table.cell(0, 3).text = buyer
        meta_table.cell(1, 0).text = "日期"
        meta_table.cell(1, 1).text = quote_date
        meta_table.cell(1, 2).text = "报价单号"
        meta_table.cell(1, 3).text = quote_no

        for r in range(2):
            for c in (0, 2):
                self._set_cell_fill(meta_table.cell(r, c), "F2F2F2")
                self._set_cell_bold(meta_table.cell(r, c))

        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        col_widths = [
            Inches(0.55),  # 序号
            Inches(2.55),  # 名称/交付物
            Inches(0.70),  # 数量
            Inches(0.65),  # 单位
            Inches(1.05),  # 单价
            Inches(1.05),  # 小计
            Inches(1.55),  # 备注
        ]
        def _apply_row_layout(cells) -> None:  # noqa: ANN001
            for idx, cell in enumerate(cells):
                cell.width = col_widths[idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        _apply_row_layout(table.rows[0].cells)

        hdr = table.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "名称/交付物"
        hdr[2].text = "数量"
        hdr[3].text = "单位"
        hdr[4].text = f"单价({currency})"
        hdr[5].text = f"小计({currency})"
        hdr[6].text = "备注"
        for cell in hdr:
            self._set_cell_fill(cell, "2F75B5")
            self._set_cell_bold(cell, color="FFFFFF")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        total = 0.0
        for i, it in enumerate(items, 1):
            name = str(it.get("name") or "")
            qty = float(it.get("quantity") or 0)
            unit = str(it.get("unit") or "项")
            unit_price = float(it.get("unit_price") or 0)
            subtotal = qty * unit_price
            total += subtotal

            row = table.add_row().cells
            _apply_row_layout(row)
            row[0].text = str(i)
            row[1].text = name
            row[2].text = _fmt_qty(qty)
            row[3].text = unit
            row[4].text = _fmt_money(unit_price)
            row[5].text = _fmt_money(subtotal)
            row[6].text = str(it.get("note") or "")

            if i % 2 == 0:
                for cell in row:
                    self._set_cell_fill(cell, "F3F6FA")

            for col_idx in (0, 2, 3, 4, 5):
                for paragraph in row[col_idx].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        total_row = table.add_row().cells
        _apply_row_layout(total_row)
        # merge 序号..单价 into one label cell
        total_row[0].merge(total_row[4])
        total_row[0].text = "合计"
        self._set_cell_fill(total_row[0], "D8E6F7")
        self._set_cell_bold(total_row[0])
        for paragraph in total_row[0].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        total_row[5].text = _fmt_money(total)
        self._set_cell_fill(total_row[5], "D8E6F7")
        self._set_cell_bold(total_row[5], color="D9534F")
        for paragraph in total_row[5].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._set_cell_fill(total_row[6], "D8E6F7")

        if note:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run("备注：")
            run.bold = True
            p.add_run(str(note))

        sign = doc.add_paragraph()
        sign.paragraph_format.space_before = Pt(10)
        sign.add_run("供方（盖章）：__________________    需方（盖章）：__________________")

        filename = f"{uuid4().hex}.docx"
        path = (self._settings.outputs_dir / filename).resolve()
        doc.save(str(path))

        file_id = filename
        token = create_download_token(settings=self._settings, file_id=file_id)
        return {"file_id": file_id, "filename": filename, "download_url": abs_url(self._settings, f"/api/files/{file_id}?token={token}")}

    async def create_quote_xlsx(
        self,
        *,
        seller: str,
        buyer: str,
        currency: str,
        items: list[dict],
        note: str | None,
    ) -> dict:
        wb = Workbook()
        wb.remove(wb.active)

        # Styles (参考 create_professional_quotation.py 的配色/字体)
        header_font = Font(name="微软雅黑", size=14, bold=True, color="FFFFFF")
        subheader_font = Font(name="微软雅黑", size=12, bold=True, color="333333")
        normal_font = Font(name="微软雅黑", size=11, color="333333")
        title_font = Font(name="微软雅黑", size=20, bold=True, color="333333")
        company_font = Font(name="微软雅黑", size=16, bold=True, color="1F497D")

        header_fill = PatternFill(start_color="2F75B5", end_color="2F75B5", fill_type="solid")
        subheader_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
        total_fill = PatternFill(start_color="D8E6F7", end_color="D8E6F7", fill_type="solid")

        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )
        thick_border = Border(
            left=Side(style="medium"),
            right=Side(style="medium"),
            top=Side(style="medium"),
            bottom=Side(style="medium"),
        )

        center = Alignment(horizontal="center", vertical="center")
        left = Alignment(horizontal="left", vertical="center", wrap_text=True)

        quote_no = f"Q-{uuid4().hex[:8].upper()}"

        # Sheet: 封面
        cover = wb.create_sheet("封面")
        cover.sheet_view.showGridLines = False
        cover["A1"] = seller
        cover["A1"].font = company_font
        cover["A1"].alignment = center
        cover.merge_cells("A1:D1")

        cover["A3"] = "报价单"
        cover["A3"].font = title_font
        cover["A3"].alignment = center
        cover.merge_cells("A3:D3")

        from datetime import datetime

        cover["A5"] = f"报价日期：{datetime.now().strftime('%Y年%m月%d日')}"
        cover["A5"].font = normal_font
        cover["C5"] = f"报价单号：{quote_no}"
        cover["C5"].font = normal_font

        cover["A7"] = "致："
        cover["A7"].font = subheader_font
        cover["B7"] = buyer
        cover["B7"].font = normal_font
        cover.merge_cells("B7:D7")

        cover["A12"] = "项目基本信息"
        cover["A12"].font = subheader_font
        cover["A12"].fill = subheader_fill
        cover["A12"].border = thin_border
        cover["A12"].alignment = left
        cover.merge_cells("A12:D12")

        total = 0.0
        for it in items:
            qty = float(it.get("quantity") or 0)
            unit_price = float(it.get("unit_price") or 0)
            total += qty * unit_price

        info_rows = [
            ("供方", seller),
            ("需方", buyer),
            ("币种", currency),
            ("合计金额", f"{total:.2f} {currency}"),
        ]
        if note:
            info_rows.append(("备注", note))

        start_row = 13
        for idx, (k, v) in enumerate(info_rows):
            r = start_row + idx
            cover[f"A{r}"] = k
            cover[f"B{r}"] = v
            cover[f"A{r}"].font = normal_font
            cover[f"B{r}"].font = normal_font
            cover[f"A{r}"].fill = PatternFill(start_color="FFFFFF", end_color="F2F2F2", fill_type="solid")
            cover[f"B{r}"].fill = PatternFill(start_color="FFFFFF", end_color="F2F2F2", fill_type="solid")
            cover[f"A{r}"].border = thin_border
            cover[f"B{r}"].border = thin_border
            cover[f"A{r}"].alignment = left
            cover[f"B{r}"].alignment = left
            cover.merge_cells(f"B{r}:D{r}")

        # Column widths (cover)
        cover.column_dimensions["A"].width = 14
        cover.column_dimensions["B"].width = 26
        cover.column_dimensions["C"].width = 18
        cover.column_dimensions["D"].width = 18

        # Sheet: 报价明细
        detail = wb.create_sheet("报价明细")
        detail.sheet_view.showGridLines = False
        detail["A1"] = "报价明细"
        detail["A1"].font = header_font
        detail["A1"].fill = header_fill
        detail["A1"].alignment = center
        detail["A1"].border = thick_border
        detail.merge_cells("A1:G1")

        headers = ["序号", "名称/交付物", "数量", "单位", f"单价({currency})", f"小计({currency})", "备注"]
        for col, h in enumerate(headers, 1):
            cell = detail[f"{get_column_letter(col)}2"]
            cell.value = h
            cell.font = subheader_font
            cell.fill = subheader_fill
            cell.border = thick_border
            cell.alignment = center

        row = 3
        total = 0.0
        for i, it in enumerate(items, 1):
            name = str(it.get("name") or "")
            qty = float(it.get("quantity") or 0)
            unit = str(it.get("unit") or "项")
            unit_price = float(it.get("unit_price") or 0)
            subtotal = qty * unit_price
            total += subtotal

            values = [i, name, qty, unit, unit_price, subtotal, str(it.get("note") or "")]
            for col, v in enumerate(values, 1):
                cell = detail[f"{get_column_letter(col)}{row}"]
                cell.value = v
                cell.font = normal_font
                cell.border = thin_border
                cell.alignment = center if col in {1, 3, 4, 5, 6} else left
                if col in {5, 6}:
                    cell.number_format = "#,##0.00"
                if col == 3:
                    cell.number_format = "0.##"
            row += 1

        total_row = row
        detail[f"A{total_row}"] = "合计"
        detail[f"A{total_row}"].font = subheader_font
        detail[f"A{total_row}"].fill = total_fill
        detail[f"A{total_row}"].border = thick_border
        detail[f"A{total_row}"].alignment = center
        detail.merge_cells(f"A{total_row}:E{total_row}")

        detail[f"F{total_row}"] = float(f"{total:.2f}")
        detail[f"F{total_row}"].font = Font(name="微软雅黑", size=12, bold=True, color="D9534F")
        detail[f"F{total_row}"].fill = total_fill
        detail[f"F{total_row}"].border = thick_border
        detail[f"F{total_row}"].alignment = center
        detail[f"F{total_row}"].number_format = "#,##0.00"

        detail[f"G{total_row}"].fill = total_fill
        detail[f"G{total_row}"].border = thick_border

        if note:
            note_row = total_row + 2
            detail[f"A{note_row}"] = "备注："
            detail[f"A{note_row}"].font = subheader_font
            detail[f"B{note_row}"] = note
            detail[f"B{note_row}"].font = normal_font
            detail[f"B{note_row}"].alignment = left
            detail.merge_cells(f"B{note_row}:G{note_row}")

        detail.freeze_panes = "A3"

        # Column widths (detail)
        widths = [8, 34, 10, 8, 14, 14, 22]
        for idx, w in enumerate(widths, 1):
            detail.column_dimensions[get_column_letter(idx)].width = w

        detail.row_dimensions[1].height = 28
        detail.row_dimensions[2].height = 20

        filename = f"{uuid4().hex}.xlsx"
        path = (self._settings.outputs_dir / filename).resolve()
        wb.save(str(path))

        file_id = filename
        token = create_download_token(settings=self._settings, file_id=file_id)
        return {"file_id": file_id, "filename": filename, "download_url": abs_url(self._settings, f"/api/files/{file_id}?token={token}")}
