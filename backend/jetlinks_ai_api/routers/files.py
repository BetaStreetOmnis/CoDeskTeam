from __future__ import annotations

import html
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from uuid import uuid4

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from fastapi.responses import FileResponse
from docx import Document
from openpyxl import load_workbook

from ..db import fetchone, row_to_dict, utc_now_iso
from ..deps import get_current_user, get_db, get_settings
from ..services.auth_service import create_download_token
from ..services.auth_service import validate_download_token
from ..url_utils import abs_url
from ..env_utils import env_str


router = APIRouter(tags=["files"])

_FILE_ID_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._-]{0,199}$")
_MIME_BY_EXT = {
    ".pptx": "application/vnd.openxmlformats-officedocument.presentationml.presentation",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    ".zip": "application/zip",
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".csv": "text/csv",
    ".json": "application/json",
    ".yaml": "application/yaml",
    ".yml": "application/yaml",
    ".log": "text/plain",
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
    ".html": "text/html; charset=utf-8",
}

_UPLOAD_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
_MAX_IMAGE_BYTES = 10 * 1024 * 1024
_MAX_FILE_BYTES = 25 * 1024 * 1024


def _safe_ext(filename: str) -> str:
    ext = Path(filename).suffix.lower() if filename else ""
    if not ext:
        return ""
    safe = "".join(ch for ch in ext if ch.isalnum() or ch in {"." , "_", "-"})
    if not safe.startswith("."):
        return ""
    if safe in {"."}:
        return ""
    if len(safe) > 16:
        return safe[:16]
    return safe


def _resolve_soffice_bin() -> str | None:
    configured = env_str("SOFFICE_BIN", "") or ""
    candidates = [
        configured,
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "soffice",
        "libreoffice",
    ]
    seen: set[str] = set()
    for raw in candidates:
        value = str(raw or "").strip()
        if not value or value in seen:
            continue
        seen.add(value)

        if os.path.isabs(value):
            path = Path(value).expanduser()
            if path.exists() and path.is_file():
                return str(path)
            continue

        resolved = shutil.which(value)
        if resolved:
            return resolved
    return None


def _resolve_pdftoppm_bin() -> str | None:
    configured = env_str("PDFTOPPM_BIN", "") or ""
    candidates = [configured, "pdftoppm"]
    seen: set[str] = set()
    for raw in candidates:
        value = str(raw or "").strip()
        if not value or value in seen:
            continue
        seen.add(value)

        if os.path.isabs(value):
            path = Path(value).expanduser()
            if path.exists() and path.is_file():
                return str(path)
            continue

        resolved = shutil.which(value)
        if resolved:
            return resolved
    return None


def _build_soffice_env(tmp_dir: Path) -> tuple[dict, str]:
    profile_dir = (tmp_dir / "lo_profile").resolve()
    profile_dir.mkdir(parents=True, exist_ok=True)
    env = os.environ.copy()
    for key in ("HOME", "USERPROFILE", "TMPDIR", "TEMP", "TMP"):
        env[key] = str(tmp_dir)
    env.setdefault("LANG", "en_US.UTF-8")
    env.setdefault("LC_ALL", "en_US.UTF-8")
    return env, f"-env:UserInstallation={profile_dir.as_uri()}"


def _soffice_convert(*, soffice: str, tmp_dir: Path, out_dir: Path, input_path: Path, fmt: str) -> None:
    env, user_install = _build_soffice_env(tmp_dir)
    cmd = [
        soffice,
        "--headless",
        "--nologo",
        "--nolockcheck",
        "--nofirststartwizard",
        "--invisible",
        user_install,
        "--convert-to",
        fmt,
        "--outdir",
        str(out_dir),
        str(input_path),
    ]
    subprocess.run(cmd, capture_output=True, check=False, timeout=75, env=env)


def _create_ppt_preview_images(*, settings, ppt_path: Path, max_slides: int = 6) -> list[str]:  # noqa: ANN001
    soffice = _resolve_soffice_bin()
    pdftoppm = _resolve_pdftoppm_bin()
    if not soffice:
        return []

    tmp_dir = Path(tempfile.mkdtemp(prefix="jetlinks-ai-ppt-preview-"))
    try:
        pngs: list[Path] = []
        pdf_path = tmp_dir / f"{ppt_path.stem}.pdf"
        _soffice_convert(soffice=soffice, tmp_dir=tmp_dir, out_dir=tmp_dir, input_path=ppt_path, fmt="pdf")

        if pdf_path.exists() and pdftoppm:
            prefix = tmp_dir / "slide"
            cmd_ppm = [
                pdftoppm,
                "-png",
                "-f",
                "1",
                "-l",
                str(max(1, int(max_slides))),
                str(pdf_path),
                str(prefix),
            ]
            subprocess.run(cmd_ppm, capture_output=True, check=False, timeout=60)
            pngs = sorted(tmp_dir.glob("slide-*.png"), key=lambda p: p.name)

        if not pngs:
            _soffice_convert(
                soffice=soffice,
                tmp_dir=tmp_dir,
                out_dir=tmp_dir,
                input_path=ppt_path,
                fmt="png:impress_png_Export",
            )
            pngs = sorted(tmp_dir.glob("*.png"), key=lambda p: p.name)

        if not pngs:
            return []

        out: list[str] = []
        for p in pngs[: max(1, int(max_slides))]:
            file_id = f"{uuid4().hex}.png"
            dest = (settings.outputs_dir / file_id).resolve()
            dest.write_bytes(p.read_bytes())
            token = create_download_token(settings=settings, file_id=file_id)
            out.append(abs_url(settings, f"/api/files/{file_id}?token={token}"))
        return out
    except Exception:
        return []
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _create_pdf_preview_images(*, settings, pdf_path: Path, max_pages: int = 6) -> list[str]:  # noqa: ANN001
    pdftoppm = _resolve_pdftoppm_bin()
    if not pdftoppm:
        return []
    if not pdf_path.exists():
        return []

    tmp_dir = Path(tempfile.mkdtemp(prefix="jetlinks-ai-pdf-preview-"))
    try:
        prefix = tmp_dir / "page"
        cmd = [
            pdftoppm,
            "-png",
            "-f",
            "1",
            "-l",
            str(max(1, int(max_pages))),
            str(pdf_path),
            str(prefix),
        ]
        subprocess.run(cmd, capture_output=True, check=False, timeout=60)

        pngs = sorted(tmp_dir.glob("page-*.png"), key=lambda p: p.name)
        if not pngs:
            return []

        out: list[str] = []
        for p in pngs:
            file_id = f"{uuid4().hex}.png"
            dest = (settings.outputs_dir / file_id).resolve()
            dest.write_bytes(p.read_bytes())
            token = create_download_token(settings=settings, file_id=file_id)
            out.append(abs_url(settings, f"/api/files/{file_id}?token={token}"))
        return out
    except Exception:
        return []
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _create_docx_preview_images(*, settings, docx_path: Path, max_pages: int = 6) -> list[str]:  # noqa: ANN001
    soffice = _resolve_soffice_bin()
    if not soffice:
        return []

    tmp_dir = Path(tempfile.mkdtemp(prefix="jetlinks-ai-docx-preview-"))
    try:
        _soffice_convert(soffice=soffice, tmp_dir=tmp_dir, out_dir=tmp_dir, input_path=docx_path, fmt="pdf")

        pdf_path = tmp_dir / f"{docx_path.stem}.pdf"
        if not pdf_path.exists():
            pdf_candidates = sorted(tmp_dir.glob("*.pdf"), key=lambda p: p.stat().st_mtime, reverse=True)
            if not pdf_candidates:
                return []
            pdf_path = pdf_candidates[0]

        return _create_pdf_preview_images(settings=settings, pdf_path=pdf_path, max_pages=max_pages)
    except Exception:
        return []
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _create_xlsx_preview_images(*, settings, xlsx_path: Path, max_pages: int = 6) -> list[str]:  # noqa: ANN001
    """Best-effort visual preview for Excel files.

    This uses LibreOffice to convert to PDF, then pdftoppm to rasterize pages.
    """
    soffice = _resolve_soffice_bin()
    if not soffice:
        return []

    tmp_dir = Path(tempfile.mkdtemp(prefix="jetlinks-ai-xlsx-preview-"))
    try:
        _soffice_convert(soffice=soffice, tmp_dir=tmp_dir, out_dir=tmp_dir, input_path=xlsx_path, fmt="pdf")

        pdf_path = tmp_dir / f"{xlsx_path.stem}.pdf"
        if not pdf_path.exists():
            pdf_candidates = sorted(tmp_dir.glob("*.pdf"), key=lambda p: p.stat().st_mtime, reverse=True)
            if not pdf_candidates:
                return []
            pdf_path = pdf_candidates[0]

        return _create_pdf_preview_images(settings=settings, pdf_path=pdf_path, max_pages=max_pages)
    except Exception:
        return []
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def _create_xlsx_preview_html(*, settings, xlsx_path: Path, max_rows: int = 50, max_cols: int = 12) -> str | None:  # noqa: ANN001
    try:
        wb = load_workbook(filename=str(xlsx_path), data_only=True, read_only=True)
    except Exception:
        return None

    ws = wb.active
    rows = list(ws.iter_rows(min_row=1, max_row=max_rows, max_col=max_cols, values_only=True))
    if not rows:
        return None

    def esc(value: object) -> str:
        return html.escape("" if value is None else str(value), quote=True)

    header = rows[0] if rows else []
    body = rows[1:] if len(rows) > 1 else []
    truncated = (ws.max_row or 0) > max_rows or (ws.max_column or 0) > max_cols

    header_cells = "".join(f"<th>{esc(c)}</th>" for c in header)
    body_rows = []
    for row in body:
        cells = "".join(f"<td>{esc(c)}</td>" for c in row)
        body_rows.append(f"<tr>{cells}</tr>")

    hint = ""
    if truncated:
        hint = "<div class=\"hint\">仅展示前 %d 行、%d 列，已截断预览。</div>" % (max_rows, max_cols)

    html_text = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>Excel 预览</title>
  <style>
    body {{ margin: 0; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Arial, sans-serif; background: #f5f7fa; color: #111827; }}
    .wrap {{ padding: 16px; }}
    table {{ width: 100%; border-collapse: collapse; background: #fff; border: 1px solid #e5e7eb; border-radius: 10px; overflow: hidden; }}
    th, td {{ border-bottom: 1px solid #eef2f7; padding: 8px 10px; font-size: 12px; text-align: left; }}
    th {{ background: #2F75B5; font-weight: 600; color: #ffffff; }}
    tr:last-child td {{ border-bottom: 0; }}
    .hint {{ margin-top: 10px; font-size: 12px; color: #6b7280; }}
  </style>
</head>
<body>
  <div class="wrap">
    <table>
      <thead><tr>{header_cells}</tr></thead>
      <tbody>
        {''.join(body_rows)}
      </tbody>
    </table>
    {hint}
  </div>
</body>
</html>
"""

    file_id = f"{uuid4().hex}.html"
    dest = (settings.outputs_dir / file_id).resolve()
    dest.write_text(html_text, encoding="utf-8")
    return file_id


def _create_docx_preview_html(*, settings, docx_path: Path, max_paragraphs: int = 260, max_table_rows: int = 40) -> str | None:  # noqa: ANN001
    """Lightweight DOCX -> HTML preview (no external converters).

    Notes:
    - This is a best-effort textual preview (not a pixel-perfect layout renderer).
    - For full-fidelity preview, deploy LibreOffice + pdftoppm so `preview_images` works.
    """
    try:
        doc = Document(str(docx_path))
    except Exception:
        return None

    blocks: list[str] = []

    def esc(value: object) -> str:
        return html.escape("" if value is None else str(value), quote=True)

    para_count = 0
    for p in doc.paragraphs:
        if para_count >= max(1, int(max_paragraphs)):
            blocks.append('<div class="hint">已截断预览（段落过多）。</div>')
            break
        text = (p.text or "").strip()
        if not text:
            continue
        style = (getattr(getattr(p, "style", None), "name", "") or "").lower()
        if "heading" in style or "标题" in style:
            blocks.append(f"<h3>{esc(text)}</h3>")
        else:
            blocks.append(f"<p>{esc(text)}</p>")
        para_count += 1

    # Basic table preview (first few tables only).
    for ti, table in enumerate(doc.tables[:6]):
        rows_html: list[str] = []
        for ri, row in enumerate(table.rows[: max(1, int(max_table_rows))]):
            cells = []
            for cell in row.cells[:24]:
                cells.append(f"<td>{esc((cell.text or '').strip())}</td>")
            rows_html.append("<tr>" + "".join(cells) + "</tr>")
        if len(table.rows) > max_table_rows:
            rows_html.append(f'<tr><td colspan="24" class="hint">表格已截断（仅前 {max_table_rows} 行）。</td></tr>')
        blocks.append(
            "<div class=\"tableWrap\">"
            f"<div class=\"tableTitle\">表格 {ti + 1}</div>"
            "<table>"
            + "".join(rows_html)
            + "</table>"
            "</div>"
        )

    if not blocks:
        return None

    html_text = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>Word 预览</title>
  <style>
    body {{ margin: 0; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, Arial, sans-serif; background: #f5f7fa; color: #111827; }}
    .wrap {{ padding: 16px; max-width: 980px; margin: 0 auto; }}
    .doc {{ background: #fff; border: 1px solid #e5e7eb; border-radius: 12px; padding: 18px 18px; }}
    h3 {{ margin: 18px 0 8px; font-size: 16px; }}
    p {{ margin: 8px 0; line-height: 1.6; font-size: 13px; }}
    table {{ width: 100%; border-collapse: collapse; margin-top: 8px; }}
    td {{ border: 1px solid #eef2f7; padding: 6px 8px; font-size: 12px; vertical-align: top; }}
    .tableWrap {{ margin-top: 14px; }}
    .tableTitle {{ font-size: 12px; color: #6b7280; margin-bottom: 6px; }}
    .hint {{ margin-top: 10px; font-size: 12px; color: #6b7280; }}
  </style>
</head>
<body>
  <div class="wrap">
    <div class="doc">
      {''.join(blocks)}
      <div class="hint">说明：此为文本预览，最终排版以 DOCX 为准。</div>
    </div>
  </div>
</body>
</html>
"""

    file_id = f"{uuid4().hex}.html"
    dest = (settings.outputs_dir / file_id).resolve()
    dest.write_text(html_text, encoding="utf-8")
    return file_id



def _safe_outputs_path(outputs_dir, file_id: str):  # noqa: ANN001
    if not _FILE_ID_RE.match(file_id) or ".." in file_id:
        raise HTTPException(status_code=400, detail="invalid file_id")

    base = outputs_dir.resolve()
    full = (base / file_id).resolve()
    if full == base or not str(full).startswith(str(base) + os.sep):
        raise HTTPException(status_code=400, detail="invalid file_id")
    return full


@router.post("/files/upload-image")
async def upload_image(
    file: UploadFile = File(...),
    project_id: int | None = Query(default=None),
    user=Depends(get_current_user),  # noqa: ANN001
    settings=Depends(get_settings),  # noqa: ANN001
    db=Depends(get_db),  # noqa: ANN001
) -> dict:
    filename_in = str(file.filename or "").strip()
    ext = _safe_ext(filename_in)
    if ext not in _UPLOAD_IMAGE_EXTS:
        raise HTTPException(status_code=400, detail=f"仅支持上传图片：{', '.join(sorted(_UPLOAD_IMAGE_EXTS))}")

    ctype = str(file.content_type or "").lower().strip()
    if not ctype.startswith("image/"):
        raise HTTPException(status_code=400, detail="仅支持上传图片（content-type 必须为 image/*）")

    pid: int | None = None
    if project_id is not None and int(project_id) > 0:
        candidate = int(project_id)
        try:
            proj_row = await fetchone(db, "SELECT id, team_id, enabled FROM team_projects WHERE id = ?", (candidate,))
        except Exception:
            pid = None
        else:
            proj = row_to_dict(proj_row) or {}
            if not proj or int(proj.get("team_id") or 0) != int(getattr(user, "team_id", 0)):
                raise HTTPException(status_code=404, detail="项目不存在")
            if not bool(proj.get("enabled")):
                raise HTTPException(status_code=400, detail="项目已禁用")
            pid = candidate

    data = await file.read()
    await file.close()

    if not data:
        raise HTTPException(status_code=400, detail="空文件")
    if len(data) > _MAX_IMAGE_BYTES:
        raise HTTPException(status_code=400, detail=f"图片太大（最大 {_MAX_IMAGE_BYTES // (1024 * 1024)}MB）")

    settings.outputs_dir.mkdir(parents=True, exist_ok=True)
    file_id = f"{uuid4().hex}{ext}"
    path = _safe_outputs_path(settings.outputs_dir, file_id)
    path.write_bytes(data)

    token = create_download_token(settings=settings, file_id=file_id)
    download_url = abs_url(settings, f"/api/files/{file_id}?token={token}")

    try:
        await db.execute(
            """
            INSERT OR IGNORE INTO file_records(
              file_id, team_id, user_id, project_id, session_id, kind, filename, content_type, size_bytes, created_at
            ) VALUES (?, ?, ?, ?, NULL, ?, ?, ?, ?, ?)
            """,
            (
                file_id,
                int(getattr(user, "team_id", 0)),
                int(getattr(user, "id", 0)),
                pid,
                "image",
                filename_in or file_id,
                ctype,
                len(data),
                utc_now_iso(),
            ),
        )
        await db.commit()
    except Exception:
        pass
    return {
        "file_id": file_id,
        "filename": filename_in or file_id,
        "content_type": ctype,
        "download_url": download_url,
    }


@router.post("/files/upload-file")
async def upload_file(
    file: UploadFile = File(...),
    project_id: int | None = Query(default=None),
    user=Depends(get_current_user),  # noqa: ANN001
    settings=Depends(get_settings),  # noqa: ANN001
    db=Depends(get_db),  # noqa: ANN001
) -> dict:
    filename_in = str(file.filename or "").strip()
    ext = _safe_ext(filename_in)

    pid: int | None = None
    if project_id is not None and int(project_id) > 0:
        candidate = int(project_id)
        try:
            proj_row = await fetchone(db, "SELECT id, team_id, enabled FROM team_projects WHERE id = ?", (candidate,))
        except Exception:
            pid = None
        else:
            proj = row_to_dict(proj_row) or {}
            if not proj or int(proj.get("team_id") or 0) != int(getattr(user, "team_id", 0)):
                raise HTTPException(status_code=404, detail="项目不存在")
            if not bool(proj.get("enabled")):
                raise HTTPException(status_code=400, detail="项目已禁用")
            pid = candidate

    data = await file.read()
    await file.close()

    if not data:
        raise HTTPException(status_code=400, detail="空文件")
    if len(data) > _MAX_FILE_BYTES:
        raise HTTPException(status_code=400, detail=f"文件太大（最大 {_MAX_FILE_BYTES // (1024 * 1024)}MB）")

    settings.outputs_dir.mkdir(parents=True, exist_ok=True)
    file_id = f"{uuid4().hex}{ext}" if ext else uuid4().hex
    path = _safe_outputs_path(settings.outputs_dir, file_id)
    path.write_bytes(data)

    token = create_download_token(settings=settings, file_id=file_id)
    download_url = abs_url(settings, f"/api/files/{file_id}?token={token}")
    ctype = str(file.content_type or "").lower().strip() or _MIME_BY_EXT.get(ext, "application/octet-stream")
    kind = "image" if ctype.startswith("image/") else "file"

    try:
        await db.execute(
            """
            INSERT OR IGNORE INTO file_records(
              file_id, team_id, user_id, project_id, session_id, kind, filename, content_type, size_bytes, created_at
            ) VALUES (?, ?, ?, ?, NULL, ?, ?, ?, ?, ?)
            """,
            (
                file_id,
                int(getattr(user, "team_id", 0)),
                int(getattr(user, "id", 0)),
                pid,
                kind,
                filename_in or file_id,
                ctype,
                len(data),
                utc_now_iso(),
            ),
        )
        await db.commit()
    except Exception:
        pass
    return {
        "file_id": file_id,
        "filename": filename_in or file_id,
        "content_type": ctype,
        "size_bytes": len(data),
        "download_url": download_url,
    }


@router.get("/files/{file_id}")
async def download_file(
    file_id: str,
    token: str = Query(min_length=1),
    settings=Depends(get_settings),  # noqa: ANN001
):
    try:
        validate_download_token(settings=settings, token=token, file_id=file_id)
    except Exception:
        raise HTTPException(status_code=401, detail="下载链接已失效，请重新生成") from None

    full = _safe_outputs_path(settings.outputs_dir, file_id)
    if not full.exists() or not full.is_file():
        raise HTTPException(status_code=404, detail="file not found")

    ext = full.suffix.lower()
    media_type = _MIME_BY_EXT.get(ext, "application/octet-stream")
    return FileResponse(path=full, media_type=media_type, filename=full.name)


@router.get("/files/inline/{file_id}")
async def inline_file(
    file_id: str,
    token: str = Query(min_length=1),
    settings=Depends(get_settings),  # noqa: ANN001
):
    try:
        validate_download_token(settings=settings, token=token, file_id=file_id)
    except Exception:
        raise HTTPException(status_code=401, detail="下载链接已失效，请重新生成") from None

    full = _safe_outputs_path(settings.outputs_dir, file_id)
    if not full.exists() or not full.is_file():
        raise HTTPException(status_code=404, detail="file not found")

    ext = full.suffix.lower()
    media_type = _MIME_BY_EXT.get(ext, "application/octet-stream")
    return FileResponse(path=full, media_type=media_type, filename=full.name, content_disposition_type="inline")


@router.get("/files/preview/{file_id}")
async def preview_file(
    file_id: str,
    _user=Depends(get_current_user),  # noqa: ANN001
    settings=Depends(get_settings),  # noqa: ANN001
):
    full = _safe_outputs_path(settings.outputs_dir, file_id)
    if not full.exists() or not full.is_file():
        raise HTTPException(status_code=404, detail="file not found")

    token = create_download_token(settings=settings, file_id=file_id)
    download_url = abs_url(settings, f"/api/files/{file_id}?token={token}")
    inline_url = abs_url(settings, f"/api/files/inline/{file_id}?token={token}")
    ext = full.suffix.lower()

    if ext == ".pptx":
        preview_images = _create_ppt_preview_images(settings=settings, ppt_path=full, max_slides=9)
        return {
            "kind": "pptx",
            "file_id": file_id,
            "download_url": download_url,
            "preview_images": preview_images,
        }

    if ext == ".pdf":
        preview_images = _create_pdf_preview_images(settings=settings, pdf_path=full, max_pages=9)
        return {
            "kind": "pdf",
            "file_id": file_id,
            "download_url": download_url,
            "preview_url": inline_url,
            "preview_images": preview_images,
        }

    if ext == ".docx":
        preview_images = _create_docx_preview_images(settings=settings, docx_path=full, max_pages=9)
        preview_file_id = _create_docx_preview_html(settings=settings, docx_path=full)
        preview_url = inline_url
        if preview_file_id:
            preview_token = create_download_token(settings=settings, file_id=preview_file_id)
            preview_url = abs_url(settings, f"/api/files/inline/{preview_file_id}?token={preview_token}")
        return {
            "kind": "docx",
            "file_id": file_id,
            "download_url": download_url,
            "preview_url": preview_url,
            "preview_images": preview_images,
        }

    if ext == ".xlsx":
        preview_images = _create_xlsx_preview_images(settings=settings, xlsx_path=full, max_pages=6)
        preview_file_id = _create_xlsx_preview_html(settings=settings, xlsx_path=full)
        preview_url = None
        if preview_file_id:
            preview_token = create_download_token(settings=settings, file_id=preview_file_id)
            preview_url = abs_url(settings, f"/api/files/inline/{preview_file_id}?token={preview_token}")
        return {
            "kind": "xlsx",
            "file_id": file_id,
            "download_url": download_url,
            "preview_images": preview_images,
            "preview_url": preview_url,
        }

    if ext == ".html":
        return {
            "kind": "html",
            "file_id": file_id,
            "download_url": download_url,
            "preview_url": inline_url,
        }

    if ext == ".zip":
        bundle_id = Path(file_id).stem
        preview_id = f"prototype_{bundle_id}"
        preview_dir = (settings.outputs_dir / preview_id).resolve()
        index_path = preview_dir / "index.html"
        if preview_dir.exists() and index_path.exists():
            preview_token = create_download_token(settings=settings, file_id=preview_id)
            preview_url = abs_url(settings, f"/api/prototype/preview/{preview_id}/index.html?token={preview_token}")
            return {
                "kind": "proto",
                "file_id": file_id,
                "download_url": download_url,
                "preview_url": preview_url,
            }

    return {"kind": "file", "file_id": file_id, "download_url": download_url}
