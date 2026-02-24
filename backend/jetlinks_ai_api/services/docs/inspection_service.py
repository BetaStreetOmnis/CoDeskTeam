from __future__ import annotations

from datetime import datetime
from uuid import uuid4

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor as DocxRGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

from ...config import Settings
from ...output_cleanup import maybe_cleanup_outputs_dir
from ...url_utils import abs_url
from ..auth_service import create_download_token


class InspectionDocService:
    def __init__(self, settings: Settings) -> None:
        self._settings = settings
        self._settings.outputs_dir.mkdir(parents=True, exist_ok=True)
        maybe_cleanup_outputs_dir(
            self._settings.outputs_dir,
            ttl_seconds=max(0, int(self._settings.outputs_ttl_hours)) * 3600,
        )

    def _apply_doc_theme(self, doc: Document) -> None:
        font_name = "微软雅黑"
        base = doc.styles["Normal"]
        base.font.name = font_name
        base.font.size = Pt(11)
        base._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

        for style_name in ("Heading 1", "Heading 2", "Heading 3"):
            if style_name in doc.styles:
                style = doc.styles[style_name]
                style.font.name = font_name
                style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _set_cell_fill(self, cell, fill: str) -> None:  # noqa: ANN001
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _set_cell_bold(self, cell, *, color: str | None = None) -> None:  # noqa: ANN001
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                if color:
                    run.font.color.rgb = DocxRGBColor.from_string(color)
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                run.font.size = Pt(11)

    async def create_inspection_docx(
        self,
        *,
        title: str,
        basic_info: dict,
        device_info: dict,
        network_info: dict,
        inspection_info: dict,
        inspection_items: list[dict],
        conclusion: dict,
        signatures: dict,
        attachments: list[str] | None = None,
    ) -> dict:
        def _v(value: object) -> str:
            s = str(value or "").strip()
            return s

        def _add_section_heading(text: str) -> None:
            table = doc.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.rows[0].cells[0]
            cell.text = text
            self._set_cell_fill(cell, "2F75B5")
            self._set_cell_bold(cell, color="FFFFFF")

        def _add_kv_table(rows: list[tuple[str, str]]) -> None:
            if not rows:
                return
            table = doc.add_table(rows=len(rows), cols=2)
            table.style = "Table Grid"
            for i, (k, v) in enumerate(rows):
                table.cell(i, 0).text = k
                table.cell(i, 1).text = v
                if i % 2 == 0:
                    self._set_cell_fill(table.cell(i, 0), "F3F6FA")
                    self._set_cell_fill(table.cell(i, 1), "F3F6FA")

        def _add_items_table(items: list[dict]) -> None:
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "序号"
            hdr[1].text = "检验项目"
            hdr[2].text = "判定（勾选）"
            hdr[3].text = "备注"
            for cell in hdr:
                self._set_cell_fill(cell, "2F75B5")
                self._set_cell_bold(cell, color="FFFFFF")
            for idx, it in enumerate(items, 1):
                name = _v(it.get("name"))
                result = _v(it.get("result")) or "□合格  □不合格  □NA"
                remark = _v(it.get("remark"))
                row = table.add_row().cells
                row[0].text = str(idx)
                row[1].text = name
                row[2].text = result
                row[3].text = remark

        doc = Document()
        self._apply_doc_theme(doc)
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(_v(title) or "物联网设备报检单")
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_run.font.name = "微软雅黑"
        title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        # 基本信息
        _add_section_heading("一、基本信息")
        report_no = _v(basic_info.get("report_no"))
        report_date = _v(basic_info.get("report_date")) or datetime.now().strftime("%Y-%m-%d")
        _add_kv_table(
            [
                ("报检编号", report_no),
                ("报检日期", report_date),
                ("报检单位（公司/部门）", _v(basic_info.get("applicant_unit"))),
                ("联系人", _v(basic_info.get("contact_name"))),
                ("电话", _v(basic_info.get("contact_phone"))),
                ("邮箱", _v(basic_info.get("contact_email"))),
                ("项目/订单号", _v(basic_info.get("project_order_no"))),
                ("报检类型", _v(basic_info.get("inspection_type"))),
                ("检验方式", _v(basic_info.get("inspection_method"))),
                ("抽样标准/比例", _v(basic_info.get("sampling_standard"))),
                ("允收标准", _v(basic_info.get("acceptance_standard"))),
            ]
        )

        # 设备信息
        _add_section_heading("二、设备信息")
        _add_kv_table(
            [
                ("设备名称", _v(device_info.get("device_name"))),
                ("型号/规格", _v(device_info.get("model"))),
                ("设备类别", _v(device_info.get("category"))),
                ("数量", _v(device_info.get("quantity"))),
                ("单位", _v(device_info.get("unit"))),
                ("序列号范围（SN）", _v(device_info.get("sn_range"))),
                ("固件/软件版本", _v(device_info.get("firmware_version"))),
                ("硬件版本", _v(device_info.get("hardware_version"))),
                ("生产厂家", _v(device_info.get("manufacturer"))),
                ("品牌", _v(device_info.get("brand"))),
                ("生产批次/日期", _v(device_info.get("batch"))),
                ("供货单位", _v(device_info.get("supplier"))),
                ("装箱单/发票/合同号", _v(device_info.get("doc_no"))),
            ]
        )

        # 网络与合规信息
        _add_section_heading("三、网络与合规信息（按需填写）")
        _add_kv_table(
            [
                ("通信方式", _v(network_info.get("communication"))),
                ("运营商/频段/发射功率", _v(network_info.get("carrier_band_power"))),
                ("SIM/ICCID", _v(network_info.get("sim_iccid"))),
                ("MAC/IMEI", _v(network_info.get("mac_imei"))),
                ("加密/安全", _v(network_info.get("security"))),
                ("设备证书/密钥载入方式", _v(network_info.get("key_provisioning"))),
                ("电源", _v(network_info.get("power"))),
                ("电池类型/容量", _v(network_info.get("battery"))),
                ("认证/符合性", _v(network_info.get("certifications"))),
                ("资料清单", _v(network_info.get("materials"))),
            ]
        )

        # 检验依据与项目
        _add_section_heading("四、检验依据与项目")
        _add_kv_table(
            [
                ("检验依据（标准/规范/图纸/协议）", _v(inspection_info.get("reference"))),
                ("检验环境（温度/湿度/供电）", _v(inspection_info.get("environment"))),
            ]
        )
        doc.add_paragraph("检验项目：")
        _add_items_table(inspection_items or [])

        # 结论
        _add_section_heading("五、检验结论")
        _add_kv_table(
            [
                ("结论", _v(conclusion.get("result"))),
                ("不合格描述/缺陷等级", _v(conclusion.get("defect_level"))),
                ("处置措施与责任方", _v(conclusion.get("disposition"))),
                ("复检要求与截止日期", _v(conclusion.get("recheck"))),
            ]
        )

        # 签字确认
        _add_section_heading("六、签字确认")
        _add_kv_table(
            [
                ("报检人", _v(signatures.get("reporter"))),
                ("检验员", _v(signatures.get("inspector"))),
                ("审核/负责人", _v(signatures.get("reviewer"))),
                ("日期", _v(signatures.get("date"))),
            ]
        )

        if attachments:
            doc.add_heading("附件清单", level=1)
            for a in attachments:
                t = _v(a)
                if t:
                    doc.add_paragraph(t)

        filename = f"{uuid4().hex}.docx"
        path = (self._settings.outputs_dir / filename).resolve()
        doc.save(str(path))

        file_id = filename
        token = create_download_token(settings=self._settings, file_id=file_id)
        return {"file_id": file_id, "filename": filename, "download_url": abs_url(self._settings, f"/api/files/{file_id}?token={token}")}

    async def create_inspection_xlsx(
        self,
        *,
        title: str,
        basic_info: dict,
        device_info: dict,
        network_info: dict,
        inspection_info: dict,
        inspection_items: list[dict],
        conclusion: dict,
        signatures: dict,
        attachments: list[str] | None = None,
    ) -> dict:
        def _v(value: object) -> str:
            return str(value or "").strip()

        wb = Workbook()
        ws = wb.active
        ws.title = "报检单"

        header_font = Font(name="微软雅黑", size=16, bold=True, color="FFFFFF")
        section_font = Font(name="微软雅黑", size=12, bold=True, color="333333")
        label_font = Font(name="微软雅黑", size=11, bold=True, color="333333")
        normal_font = Font(name="微软雅黑", size=11, color="333333")

        header_fill = PatternFill(start_color="2F75B5", end_color="2F75B5", fill_type="solid")
        section_fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")

        thin_border = Border(
            left=Side(style="thin"),
            right=Side(style="thin"),
            top=Side(style="thin"),
            bottom=Side(style="thin"),
        )

        left = Alignment(horizontal="left", vertical="center", wrap_text=True)
        center = Alignment(horizontal="center", vertical="center", wrap_text=True)

        def _merge_row(row: int, text: str, *, fill) -> None:  # noqa: ANN001
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
            c = ws.cell(row=row, column=1, value=text)
            c.font = section_font if fill == section_fill else header_font
            c.fill = fill
            c.alignment = center
            for col in range(1, 7):
                cc = ws.cell(row=row, column=col)
                cc.border = thin_border
                cc.fill = fill

        def _kv(row: int, key: str, value: str) -> int:
            ws.cell(row=row, column=1, value=key).font = label_font
            ws.cell(row=row, column=1).alignment = left
            ws.cell(row=row, column=1).border = thin_border

            ws.merge_cells(start_row=row, start_column=2, end_row=row, end_column=6)
            vcell = ws.cell(row=row, column=2, value=value)
            vcell.font = normal_font
            vcell.alignment = left
            for col in range(2, 7):
                ws.cell(row=row, column=col).border = thin_border
            return row + 1

        r = 1
        _merge_row(r, _v(title) or "物联网设备报检单", fill=header_fill)
        ws.row_dimensions[r].height = 28
        r += 2

        _merge_row(r, "一、基本信息", fill=section_fill)
        r += 1
        r = _kv(r, "报检编号", _v(basic_info.get("report_no")))
        r = _kv(r, "报检日期", _v(basic_info.get("report_date")) or datetime.now().strftime("%Y-%m-%d"))
        r = _kv(r, "报检单位（公司/部门）", _v(basic_info.get("applicant_unit")))
        r = _kv(r, "联系人", _v(basic_info.get("contact_name")))
        r = _kv(r, "电话", _v(basic_info.get("contact_phone")))
        r = _kv(r, "邮箱", _v(basic_info.get("contact_email")))
        r = _kv(r, "项目/订单号", _v(basic_info.get("project_order_no")))
        r = _kv(r, "报检类型", _v(basic_info.get("inspection_type")))
        r = _kv(r, "检验方式", _v(basic_info.get("inspection_method")))
        r = _kv(r, "抽样标准/比例", _v(basic_info.get("sampling_standard")))
        r = _kv(r, "允收标准", _v(basic_info.get("acceptance_standard")))
        r += 1

        _merge_row(r, "二、设备信息", fill=section_fill)
        r += 1
        r = _kv(r, "设备名称", _v(device_info.get("device_name")))
        r = _kv(r, "型号/规格", _v(device_info.get("model")))
        r = _kv(r, "设备类别", _v(device_info.get("category")))
        r = _kv(r, "数量/单位", f"{_v(device_info.get('quantity'))} {_v(device_info.get('unit'))}".strip())
        r = _kv(r, "序列号范围（SN）", _v(device_info.get("sn_range")))
        r = _kv(r, "固件/软件版本", _v(device_info.get("firmware_version")))
        r = _kv(r, "硬件版本", _v(device_info.get("hardware_version")))
        r = _kv(r, "生产厂家/品牌", f"{_v(device_info.get('manufacturer'))} {_v(device_info.get('brand'))}".strip())
        r = _kv(r, "生产批次/日期", _v(device_info.get("batch")))
        r = _kv(r, "供货单位", _v(device_info.get("supplier")))
        r = _kv(r, "装箱单/发票/合同号", _v(device_info.get("doc_no")))
        r += 1

        _merge_row(r, "三、网络与合规信息（按需填写）", fill=section_fill)
        r += 1
        r = _kv(r, "通信方式", _v(network_info.get("communication")))
        r = _kv(r, "运营商/频段/发射功率", _v(network_info.get("carrier_band_power")))
        r = _kv(r, "SIM/ICCID", _v(network_info.get("sim_iccid")))
        r = _kv(r, "MAC/IMEI", _v(network_info.get("mac_imei")))
        r = _kv(r, "加密/安全", _v(network_info.get("security")))
        r = _kv(r, "证书/密钥载入方式", _v(network_info.get("key_provisioning")))
        r = _kv(r, "电源/电池", f"{_v(network_info.get('power'))} {_v(network_info.get('battery'))}".strip())
        r = _kv(r, "认证/符合性", _v(network_info.get("certifications")))
        r = _kv(r, "资料清单", _v(network_info.get("materials")))
        r += 1

        _merge_row(r, "四、检验依据与项目", fill=section_fill)
        r += 1
        r = _kv(r, "检验依据（标准/规范/图纸/协议）", _v(inspection_info.get("reference")))
        r = _kv(r, "检验环境（温度/湿度/供电）", _v(inspection_info.get("environment")))
        r += 1

        # Items table header
        ws.cell(row=r, column=1, value="序号").font = section_font
        ws.cell(row=r, column=2, value="检验项目").font = section_font
        ws.cell(row=r, column=3, value="判定（勾选）").font = section_font
        ws.cell(row=r, column=4, value="备注").font = section_font
        for col in range(1, 5):
            c = ws.cell(row=r, column=col)
            c.fill = section_fill
            c.border = thin_border
            c.alignment = center
        ws.merge_cells(start_row=r, start_column=4, end_row=r, end_column=6)
        for col in range(4, 7):
            ws.cell(row=r, column=col).fill = section_fill
            ws.cell(row=r, column=col).border = thin_border
            ws.cell(row=r, column=col).alignment = center
        r += 1

        for idx, it in enumerate(inspection_items or [], 1):
            ws.cell(row=r, column=1, value=idx).font = normal_font
            ws.cell(row=r, column=2, value=_v(it.get("name"))).font = normal_font
            ws.cell(row=r, column=3, value=_v(it.get("result")) or "□合格 □不合格 □NA").font = normal_font
            ws.merge_cells(start_row=r, start_column=4, end_row=r, end_column=6)
            ws.cell(row=r, column=4, value=_v(it.get("remark"))).font = normal_font
            for col in range(1, 7):
                ws.cell(row=r, column=col).border = thin_border
                ws.cell(row=r, column=col).alignment = left if col in {2, 4} else center
            r += 1
        r += 1

        _merge_row(r, "五、检验结论", fill=section_fill)
        r += 1
        r = _kv(r, "结论", _v(conclusion.get("result")))
        r = _kv(r, "不合格描述/缺陷等级", _v(conclusion.get("defect_level")))
        r = _kv(r, "处置措施与责任方", _v(conclusion.get("disposition")))
        r = _kv(r, "复检要求与截止日期", _v(conclusion.get("recheck")))
        r += 1

        _merge_row(r, "六、签字确认", fill=section_fill)
        r += 1
        r = _kv(r, "报检人", _v(signatures.get("reporter")))
        r = _kv(r, "检验员", _v(signatures.get("inspector")))
        r = _kv(r, "审核/负责人", _v(signatures.get("reviewer")))
        r = _kv(r, "日期", _v(signatures.get("date")))
        r += 1

        if attachments:
            _merge_row(r, "附件清单", fill=section_fill)
            r += 1
            for a in attachments:
                t = _v(a)
                if not t:
                    continue
                r = _kv(r, "附件", t)

        # Column widths
        ws.column_dimensions["A"].width = 14
        ws.column_dimensions["B"].width = 28
        ws.column_dimensions["C"].width = 18
        ws.column_dimensions["D"].width = 16
        ws.column_dimensions["E"].width = 16
        ws.column_dimensions["F"].width = 16

        filename = f"{uuid4().hex}.xlsx"
        path = (self._settings.outputs_dir / filename).resolve()
        wb.save(str(path))

        file_id = filename
        token = create_download_token(settings=self._settings, file_id=file_id)
        return {"file_id": file_id, "filename": filename, "download_url": abs_url(self._settings, f"/api/files/{file_id}?token={token}")}
