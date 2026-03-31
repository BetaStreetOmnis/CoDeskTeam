from __future__ import annotations
import asyncio
import base64
import csv
import io
import json
import os
import re
import shutil
import subprocess
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any
import httpx
from docx import Document
from openpyxl import Workbook, load_workbook
from PIL import Image, ImageDraw, ImageEnhance, ImageFilter, ImageFont
from uuid import uuid4
from ..agent.providers.openai_provider import OpenAiProvider
from ..agent.types import ChatMessage
from ..config import Settings
from ..db import fetchone, row_to_dict, utc_now_iso
from ..services.auth_service import create_download_token
from ..url_utils import abs_url
from .doc_service import DocService
from .prototype_service import PrototypeService
_FILE_ID_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._-]{0,199}$")
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff", ".avif"}
_VIDEO_EXTS = {".mp4", ".mov", ".mkv", ".avi", ".m4v", ".webm", ".ts"}
_AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".aac", ".flac", ".ogg", ".opus"}
_TEXT_EXTS = {".txt", ".md", ".csv", ".json"}
_SPREADSHEET_EXTS = {".xlsx"}
_MODE_SET = {"hybrid", "cloud", "edge"}
_VISION_OPERATION_SET = {"auto", "generate", "edit"}
_ALLOWED_IMAGE_SIZES = {"1024x1024", "1536x1024", "1024x1536"}
_ALLOWED_IMAGE_QUALITIES = {"auto", "low", "medium", "high"}
_ALLOWED_IMAGE_BACKGROUNDS = {"auto", "transparent", "opaque"}
_ALLOWED_IMAGE_OUTPUTS = {"png", "jpeg", "webp", "avif"}
_ALLOWED_INPUT_FIDELITIES = {"low", "high"}
_ALLOWED_EDIT_SCOPES = {"auto", "background", "subject", "region"}
_ALLOWED_REFERENCE_STRENGTHS = {"low", "medium", "high"}
_V1_RE = re.compile(r"/v1(?:$|/)")
_CROP_RATIO_BY_PRESET = {"1:1": (1, 1), "4:3": (4, 3), "16:9": (16, 9)}
_WATERMARK_POSITIONS = {"top_left", "top_right", "bottom_left", "bottom_right", "center"}
@dataclass
class PipelineArtifact:
    file_id: str
    filename: str
    kind: str
    detail: str
    step: str
    download_url: str
    content_type: str
    size_bytes: int
    def to_dict(self) -> dict[str, Any]:
        return {
            "file_id": self.file_id,
            "filename": self.filename,
            "kind": self.kind,
            "detail": self.detail,
            "step": self.step,
            "download_url": self.download_url,
            "content_type": self.content_type,
            "size_bytes": self.size_bytes,
        }
def _safe_suffix(name: str) -> str:
    raw = str(name or "").strip()
    if raw.startswith(".") and raw.count(".") == 1 and len(raw) > 1:
        suffix = raw.lower()
    else:
        suffix = Path(raw).suffix.lower().strip()
    safe = "".join(ch for ch in suffix if ch.isalnum() or ch in {".", "-", "_"})
    if not safe.startswith("."):
        return ""
    if len(safe) > 16:
        return safe[:16]
    return safe
def _content_type_by_suffix(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".png"}:
        return "image/png"
    if suffix in {".jpg", ".jpeg"}:
        return "image/jpeg"
    if suffix == ".webp":
        return "image/webp"
    if suffix == ".avif":
        return "image/avif"
    if suffix == ".gif":
        return "image/gif"
    if suffix in {".mp4", ".m4v"}:
        return "video/mp4"
    if suffix in {".mov"}:
        return "video/quicktime"
    if suffix in {".webm"}:
        return "video/webm"
    if suffix in {".mp3"}:
        return "audio/mpeg"
    if suffix in {".wav"}:
        return "audio/wav"
    if suffix in {".m4a"}:
        return "audio/mp4"
    if suffix in {".csv"}:
        return "text/csv"
    if suffix in {".md"}:
        return "text/markdown"
    if suffix in {".json"}:
        return "application/json"
    if suffix in {".docx"}:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if suffix in {".xlsx"}:
        return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if suffix in {".pptx"}:
        return "application/vnd.openxmlformats-officedocument.presentationml.presentation"
    if suffix in {".zip"}:
        return "application/zip"
    return "application/octet-stream"
def _trim(value: Any) -> str:
    return str(value or "").strip()
def _safe_stem(value: Any, *, max_len: int = 40) -> str:
    text = _trim(value).lower()
    text = re.sub(r"[^a-z0-9._-]+", "-", text)
    text = re.sub(r"-{2,}", "-", text).strip("-._")
    if not text:
        text = "file"
    return text[:max_len]
def _normalize_mode(value: Any) -> str:
    mode = _trim(value).lower()
    if mode in _MODE_SET:
        return mode
    return "hybrid"
def _normalize_openai_base_url(base_url: str) -> str:
    base = _trim(base_url).rstrip("/")
    if not base:
        return "https://api.openai.com/v1"
    if "open.bigmodel.cn/api/paas/v4" in base.lower():
        return base
    if _V1_RE.search(base):
        return base
    return f"{base}/v1"
def _normalize_vision_operation(value: Any, *, has_inputs: bool) -> str:
    op = _trim(value).lower()
    if op in _VISION_OPERATION_SET:
        return op
    return "edit" if has_inputs else "generate"
def _normalize_image_size(value: Any) -> str:
    size = _trim(value).lower()
    if size in _ALLOWED_IMAGE_SIZES:
        return size
    return "1024x1024"
def _pick_closest_image_size(width: int, height: int) -> str:
    if width <= 0 or height <= 0:
        return "1024x1024"
    ratio = width / max(1, height)
    candidates = {
        "1024x1024": 1.0,
        "1536x1024": 1.5,
        "1024x1536": 2 / 3,
    }
    return min(candidates.items(), key=lambda item: abs(item[1] - ratio))[0]
def _normalize_image_quality(value: Any) -> str:
    quality = _trim(value).lower()
    if quality in _ALLOWED_IMAGE_QUALITIES:
        return quality
    return "auto"
def _normalize_image_background(value: Any) -> str:
    background = _trim(value).lower()
    if background in _ALLOWED_IMAGE_BACKGROUNDS:
        return background
    return "auto"
def _normalize_image_output(value: Any) -> str:
    output = _trim(value).lower()
    if output in _ALLOWED_IMAGE_OUTPUTS:
        return output
    return "png"
def _normalize_input_fidelity(value: Any) -> str:
    fidelity = _trim(value).lower()
    if fidelity in _ALLOWED_INPUT_FIDELITIES:
        return fidelity
    return "high"
def _normalize_edit_scope(value: Any) -> str:
    scope = _trim(value).lower()
    if scope in _ALLOWED_EDIT_SCOPES:
        return scope
    return "auto"
def _normalize_reference_strength(value: Any) -> str:
    strength = _trim(value).lower()
    if strength in _ALLOWED_REFERENCE_STRENGTHS:
        return strength
    return "high"
def _image_suffix_and_content_type(output_format: str) -> tuple[str, str]:
    if output_format == "jpeg":
        return ".jpg", "image/jpeg"
    if output_format == "webp":
        return ".webp", "image/webp"
    if output_format == "avif":
        return ".avif", "image/avif"
    return ".png", "image/png"
def _image_suffix_and_content_type_by_header(content_type: str | None, output_format: str) -> tuple[str, str]:
    normalized = _trim(content_type).split(";", 1)[0].lower()
    if normalized == "image/jpeg":
        return ".jpg", "image/jpeg"
    if normalized == "image/webp":
        return ".webp", "image/webp"
    if normalized == "image/avif":
        return ".avif", "image/avif"
    if normalized == "image/png":
        return ".png", "image/png"
    return _image_suffix_and_content_type(output_format)
def _infer_image_content_type_from_bytes(data: bytes) -> str | None:
    if not data:
        return None
    try:
        with Image.open(io.BytesIO(data)) as img:
            fmt = _trim(getattr(img, "format", "")).upper()
    except Exception:
        return None
    if fmt in {"JPEG", "JPG"}:
        return "image/jpeg"
    if fmt == "PNG":
        return "image/png"
    if fmt == "WEBP":
        return "image/webp"
    if fmt == "AVIF":
        return "image/avif"
    if fmt == "GIF":
        return "image/gif"
    return None
def _glm_image_quality(value: str) -> str:
    quality = _trim(value).lower()
    if quality in {"medium", "high"}:
        return "hd"
    return "standard"
def _extract_http_error_text(resp: httpx.Response) -> str:
    try:
        payload = resp.json()
    except Exception:
        payload = None
    if isinstance(payload, dict):
        detail = payload.get("error")
        if isinstance(detail, dict):
            message = _trim(detail.get("message"))
            if message:
                return message
        message = _trim(payload.get("detail"))
        if message:
            return message
    return _trim(resp.text) or f"HTTP {resp.status_code}"
def _supports_image_edit_fallback(resp: httpx.Response) -> bool:
    if int(resp.status_code) not in {404, 405, 501}:
        return False
    detail = _extract_http_error_text(resp).lower()
    if any(token in detail for token in {"incorrect api key", "unauthorized", "invalid_api_key", "authentication"}):
        return False
    return any(token in detail for token in {"cannot post", "not found", "unsupported", "images/edits"})
def _is_retryable_image_error(resp: httpx.Response) -> bool:
    if int(resp.status_code) in {429, 500, 502, 503, 504}:
        return True
    detail = _extract_http_error_text(resp).lower()
    return any(
        token in detail
        for token in {
            "访问量过大",
            "rate limit",
            "too many requests",
            "temporarily unavailable",
            "server busy",
            "overloaded",
            "try again later",
        }
    )
def _first_image_payload(data: dict[str, Any]) -> tuple[bytes | None, str | None, str | None]:
    items = data.get("data")
    if not isinstance(items, list) or not items:
        return None, None, None
    first = items[0] if isinstance(items[0], dict) else {}
    b64 = _trim(first.get("b64_json"))
    if b64:
        try:
            return base64.b64decode(b64), None, _trim(first.get("revised_prompt")) or None
        except Exception:
            return None, None, _trim(first.get("revised_prompt")) or None
    url = _trim(first.get("url"))
    if url:
        return None, url, _trim(first.get("revised_prompt")) or None
    return None, None, _trim(first.get("revised_prompt")) or None
def _coerce_processing_input(path: Path) -> dict[str, Any]:
    return {
        "file_id": path.name,
        "filename": path.name,
        "content_type": _content_type_by_suffix(path),
        "size_bytes": int(path.stat().st_size) if path.exists() else 0,
        "path": path,
    }
def _safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(value)
    except Exception:
        return default
def _should_post_process(*, enhance: bool, width: int, formats: list[str], primary_format: str) -> bool:
    normalized_formats = [fmt for fmt in formats if fmt]
    if enhance or width > 0:
        return True
    if not normalized_formats:
        return False
    if len(normalized_formats) > 1:
        return True
    return normalized_formats[0] != primary_format
def _normalize_target_formats(values: Any, *, defaults: list[str]) -> list[str]:
    raw = [str(v).strip().lower() for v in (values or defaults) if str(v).strip()]
    normalized = [fmt for fmt in raw if fmt in {"png", "jpg", "webp", "avif"}]
    return normalized or defaults
def _normalize_crop_presets(values: Any) -> list[str]:
    presets = [str(v).strip() for v in (values or []) if str(v).strip()]
    return [item for item in presets if item in _CROP_RATIO_BY_PRESET]
def _normalize_watermark_position(value: Any) -> str:
    position = _trim(value).lower()
    return position if position in _WATERMARK_POSITIONS else "bottom_right"
def _parse_datetime(value: Any) -> datetime | None:
    raw = _trim(value)
    if not raw:
        return None
    for fmt in ("%Y-%m-%d %H:%M", "%Y-%m-%dT%H:%M", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d"):
        try:
            dt = datetime.strptime(raw, fmt)
            if fmt == "%Y-%m-%d":
                dt = dt.replace(hour=9, minute=0)
            return dt
        except Exception:
            continue
    try:
        return datetime.fromisoformat(raw.replace("Z", "+00:00")).replace(tzinfo=None)
    except Exception:
        return None
def _format_dt(value: datetime | None) -> str:
    if not value:
        return "TBD"
    return value.strftime("%Y-%m-%d %H:%M")
def _escape_ics(text: str) -> str:
    return text.replace("\\", "\\\\").replace(",", "\\,").replace(";", "\\;").replace("\n", "\\n")
def _text_from_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n".join(parts).strip()
def _csv_text_to_xlsx_bytes(text: str) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    reader = csv.reader(io.StringIO(text))
    for row in reader:
        ws.append(list(row))
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
def _json_text_to_csv(text: str) -> str:
    try:
        payload = json.loads(text)
    except Exception:
        return "value\n" + text.replace("\n", " ")
    rows = payload if isinstance(payload, list) else [payload]
    dict_rows = [row for row in rows if isinstance(row, dict)]
    if not dict_rows:
        return "value\n" + json.dumps(payload, ensure_ascii=False)
    headers: list[str] = []
    for row in dict_rows:
        for key in row.keys():
            if key not in headers:
                headers.append(str(key))
    sio = io.StringIO()
    writer = csv.DictWriter(sio, fieldnames=headers)
    writer.writeheader()
    for row in dict_rows:
        writer.writerow({key: row.get(key, "") for key in headers})
    return sio.getvalue()
def _xlsx_to_csv_text(path: Path) -> str:
    wb = load_workbook(path, data_only=True)
    ws = wb.active
    sio = io.StringIO()
    writer = csv.writer(sio)
    for row in ws.iter_rows(values_only=True):
        writer.writerow(["" if cell is None else cell for cell in row])
    return sio.getvalue()
def _bytes_to_png_if_needed(image: Image.Image) -> bytes:
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return buf.getvalue()
def _fit_generated_image_to_source(generated_bytes: bytes, source_path: Path) -> tuple[bytes, str] | None:
    try:
        with Image.open(io.BytesIO(generated_bytes)) as generated_img, Image.open(source_path) as source_img:
            generated = generated_img.convert("RGBA")
            source = source_img.convert("RGBA")
            source_ratio = source.width / max(1, source.height)
            generated_ratio = generated.width / max(1, generated.height)
            if generated_ratio > source_ratio:
                crop_width = int(generated.height * source_ratio)
                left = max(0, (generated.width - crop_width) // 2)
                generated = generated.crop((left, 0, left + crop_width, generated.height))
            elif generated_ratio < source_ratio:
                crop_height = int(generated.width / source_ratio)
                top = max(0, (generated.height - crop_height) // 2)
                generated = generated.crop((0, top, generated.width, top + crop_height))
            generated = generated.resize((source.width, source.height), Image.LANCZOS)
            return _bytes_to_png_if_needed(generated), "image/png"
    except Exception:
        return None
def _build_background_mask(source: Image.Image) -> Image.Image:
    rgba = source.convert("RGBA")
    mask = Image.new("L", rgba.size, 0)
    src = rgba.load()
    dst = mask.load()
    for y in range(rgba.height):
        for x in range(rgba.width):
            r, g, b, a = src[x, y]
            if a < 32:
                dst[x, y] = 255
                continue
            max_rgb = max(r, g, b)
            min_rgb = min(r, g, b)
            if max_rgb >= 238 and (max_rgb - min_rgb) <= 20:
                dst[x, y] = 255
    return mask.filter(ImageFilter.GaussianBlur(radius=6))
def _edit_alpha_by_strength(strength: str) -> int:
    if strength == "low":
        return 224
    if strength == "medium":
        return 176
    return 136
def _region_box(size: tuple[int, int], region: dict[str, Any] | None) -> tuple[int, int, int, int]:
    width, height = size
    if not isinstance(region, dict):
        margin_x = max(1, int(width * 0.18))
        margin_y = max(1, int(height * 0.18))
        return (margin_x, margin_y, width - margin_x, height - margin_y)

    def _coord(value: Any, total: int, default: float) -> int:
        try:
            num = float(value)
        except Exception:
            num = default
        if 0 <= num <= 1:
            return int(total * num)
        return int(num)

    x = _coord(region.get("x"), width, 0.2)
    y = _coord(region.get("y"), height, 0.2)
    w = _coord(region.get("width"), width, width * 0.6)
    h = _coord(region.get("height"), height, height * 0.6)
    left = max(0, min(width - 1, x))
    top = max(0, min(height - 1, y))
    right = max(left + 1, min(width, left + max(1, w)))
    bottom = max(top + 1, min(height, top + max(1, h)))
    return left, top, right, bottom
def _composite_generated_edit(
    *,
    source_path: Path,
    generated_bytes: bytes,
    edit_scope: str,
    reference_strength: str,
    edit_region: dict[str, Any] | None,
) -> tuple[bytes, str] | None:
    try:
        with Image.open(source_path) as source_img, Image.open(io.BytesIO(generated_bytes)) as generated_img:
            source = source_img.convert("RGBA")
            generated = generated_img.convert("RGBA")
            if generated.size != source.size:
                generated = generated.resize(source.size, Image.LANCZOS)

            if edit_scope == "auto":
                return _bytes_to_png_if_needed(generated), "image/png"

            alpha = _edit_alpha_by_strength(reference_strength)
            if edit_scope == "region":
                mask = Image.new("L", source.size, 0)
                draw = ImageDraw.Draw(mask)
                left, top, right, bottom = _region_box(source.size, edit_region)
                draw.rounded_rectangle((left, top, right, bottom), radius=max(8, int(min(source.size) * 0.03)), fill=alpha)
                mask = mask.filter(ImageFilter.GaussianBlur(radius=10))
            else:
                bg_mask = _build_background_mask(source)
                mask = bg_mask if edit_scope == "background" else Image.eval(bg_mask, lambda px: 255 - px)
                if alpha != 255:
                    mask = mask.point(lambda px: int(px * (alpha / 255)))
                mask = mask.filter(ImageFilter.GaussianBlur(radius=4))

            composed = Image.composite(generated, source, mask)
            return _bytes_to_png_if_needed(composed), "image/png"
    except Exception:
        return None
def _apply_transparent_background(image: Image.Image) -> Image.Image:
    rgba = image.convert("RGBA")
    pixels = []
    for r, g, b, a in rgba.getdata():
        if r >= 245 and g >= 245 and b >= 245:
            pixels.append((r, g, b, 0))
        else:
            pixels.append((r, g, b, a))
    rgba.putdata(pixels)
    return rgba
def _apply_filter_preset(image: Image.Image, preset: str, *, enhance: bool) -> Image.Image:
    result = image.convert("RGBA")
    preset_name = _trim(preset).lower()
    if enhance or preset_name in {"clarity", "business"}:
        result = ImageEnhance.Contrast(result).enhance(1.08)
        result = ImageEnhance.Color(result).enhance(1.06)
        result = result.filter(ImageFilter.UnsharpMask(radius=1.5, percent=120, threshold=3))
    if preset_name == "warm":
        result = ImageEnhance.Color(result).enhance(1.12)
        result = ImageEnhance.Brightness(result).enhance(1.03)
    elif preset_name == "dramatic":
        result = ImageEnhance.Contrast(result).enhance(1.18)
        result = ImageEnhance.Brightness(result).enhance(0.98)
    elif preset_name == "business":
        result = ImageEnhance.Color(result).enhance(0.96)
    return result
def _resize_image(image: Image.Image, *, width: int) -> Image.Image:
    if width <= 0 or image.width <= width:
        return image
    height = max(1, int(image.height * (width / image.width)))
    return image.resize((width, height), Image.LANCZOS)
def _crop_image_to_ratio(image: Image.Image, ratio_key: str) -> Image.Image:
    ratio = _CROP_RATIO_BY_PRESET.get(ratio_key)
    if not ratio:
        return image
    rw, rh = ratio
    source_ratio = image.width / max(1, image.height)
    target_ratio = rw / rh
    if source_ratio > target_ratio:
        crop_width = int(image.height * target_ratio)
        left = max(0, (image.width - crop_width) // 2)
        return image.crop((left, 0, left + crop_width, image.height))
    crop_height = int(image.width / target_ratio)
    top = max(0, (image.height - crop_height) // 2)
    return image.crop((0, top, image.width, top + crop_height))
def _apply_watermark(image: Image.Image, *, text: str, position: str) -> Image.Image:
    if not _trim(text):
        return image
    canvas = image.convert("RGBA")
    overlay = Image.new("RGBA", canvas.size, (255, 255, 255, 0))
    draw = ImageDraw.Draw(overlay)
    font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), text, font=font)
    tw = bbox[2] - bbox[0]
    th = bbox[3] - bbox[1]
    margin = max(12, int(min(canvas.size) * 0.03))
    if position == "top_left":
        xy = (margin, margin)
    elif position == "top_right":
        xy = (max(margin, canvas.width - tw - margin), margin)
    elif position == "bottom_left":
        xy = (margin, max(margin, canvas.height - th - margin))
    elif position == "center":
        xy = (max(margin, (canvas.width - tw) // 2), max(margin, (canvas.height - th) // 2))
    else:
        xy = (max(margin, canvas.width - tw - margin), max(margin, canvas.height - th - margin))
    draw.text(xy, text, fill=(255, 255, 255, 150), font=font)
    return Image.alpha_composite(canvas, overlay)
def _language_tag(lang: str) -> str:
    normalized = _trim(lang) or "zh-CN"
    return normalized
def _offline_translate(text: str, lang: str) -> str:
    if lang.lower() in {"zh", "zh-cn", "zh-hans"}:
        return text
    return f"[{_language_tag(lang)}]\n{text}"
def _review_suggestions(text: str, focus: list[str], brand_terms: list[str]) -> list[str]:
    suggestions: list[str] = []
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return ["建议补充正文内容后再执行排版优化。"]
    if len(lines) < 3:
        suggestions.append("内容段落较少，建议补充背景、目标、方案与验收四段结构。")
    if all(len(line) > 80 for line in lines[:2]):
        suggestions.append("前两段偏长，建议增加小标题并拆分长句。")
    if "结构层级" in focus:
        suggestions.append("建议使用“目标 / 范围 / 步骤 / 风险 / 验收”五段结构。")
    if "可读性" in focus or "摘要提炼" in focus:
        suggestions.append("建议在首段增加 2~3 句摘要，便于快速阅读。")
    if brand_terms:
        missing = [term for term in brand_terms if term and term not in text]
        if missing:
            suggestions.append("建议补充品牌关键词：" + "、".join(missing[:5]))
    return suggestions[:6]
def _version_summary(source: str, revised: str) -> list[str]:
    source_len = len(source.strip())
    revised_len = len(revised.strip())
    return [
        f"原文长度：{source_len} 字符",
        f"修订长度：{revised_len} 字符",
        "已补齐标题层级与摘要结构。" if "\n## " in revised or "\n### " in revised else "结构仍偏平铺，可继续增强标题层级。",
    ]
def _build_docx_bytes(title: str, sections: list[tuple[str, str]]) -> bytes:
    doc = Document()
    doc.add_heading(title or "文档", level=0)
    for heading, body in sections:
        doc.add_heading(heading, level=1)
        for para in [line.strip() for line in body.split("\n") if line.strip()]:
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
def _safe_file_id(file_id: str) -> bool:
    return bool(_FILE_ID_RE.match(file_id or "")) and ".." not in file_id
def _ensure_outputs_path(outputs_dir: Path, file_id: str) -> Path:
    if not _safe_file_id(file_id):
        raise ValueError("invalid file_id")
    base = outputs_dir.resolve()
    full = (base / file_id).resolve()
    if full == base or not str(full).startswith(str(base) + os.sep):
        raise ValueError("invalid file path")
    return full
class ToolkitPipelineService:
    def __init__(self, settings: Settings) -> None:
        self._settings = settings
        self._doc_service = DocService(settings)
        self._prototype_service = PrototypeService(settings)
        self._settings.outputs_dir.mkdir(parents=True, exist_ok=True)
    async def run_pipeline(
        self,
        *,
        pipeline_id: str,
        payload: dict[str, Any],
        user,
        db,
        project_id: int | None = None,
    ) -> dict[str, Any]:
        normalized_id = _trim(pipeline_id).lower()
        mode = _normalize_mode(payload.get("deployment_mode"))
        if normalized_id == "vision":
            summary, artifacts, warnings = await self._run_vision_pipeline(
                payload=payload, mode=mode, user=user, db=db, project_id=project_id
            )
            name = "视觉设计与图像处理 Pipeline"
        elif normalized_id == "media":
            summary, artifacts, warnings = await self._run_media_pipeline(
                payload=payload, mode=mode, user=user, db=db, project_id=project_id
            )
            name = "视频与音频处理 Pipeline"
        elif normalized_id == "content":
            summary, artifacts, warnings = await self._run_content_pipeline(
                payload=payload, mode=mode, user=user, db=db, project_id=project_id
            )
            name = "内容创作与写作 Pipeline"
        elif normalized_id == "office":
            summary, artifacts, warnings = await self._run_office_pipeline(
                payload=payload, mode=mode, user=user, db=db, project_id=project_id
            )
            name = "办公与效率提升 Pipeline"
        elif normalized_id == "full":
            summary, artifacts, warnings = await self._run_full_pipeline(
                payload=payload, mode=mode, user=user, db=db, project_id=project_id
            )
            name = "AI+通用工具集全链路 Pipeline"
        else:
            raise ValueError("未知 pipeline_id")
        primary = artifacts[0] if artifacts else None
        return {
            "pipeline_id": normalized_id,
            "pipeline_name": name,
            "deployment_mode": mode,
            "summary": summary,
            "artifacts": [item.to_dict() for item in artifacts],
            "warnings": warnings,
            "download_url": primary.download_url if primary else "",
            "file_id": primary.file_id if primary else "",
            "filename": primary.filename if primary else "",
        }
    async def _resolve_input_files(self, *, file_ids: list[str], user, db) -> tuple[list[dict[str, Any]], list[str]]:
        resolved: list[dict[str, Any]] = []
        warnings: list[str] = []
        for raw in file_ids:
            file_id = _trim(raw)
            if not file_id:
                continue
            if not _safe_file_id(file_id):
                warnings.append(f"跳过非法 file_id: {file_id}")
                continue
            row = await fetchone(
                db,
                """
                SELECT file_id, team_id, user_id, filename, content_type, size_bytes
                FROM file_records
                WHERE file_id = ?
                """,
                (file_id,),
            )
            item = row_to_dict(row) or {}
            if not item:
                warnings.append(f"未找到输入文件记录: {file_id}")
                continue
            if int(item.get("team_id") or 0) != int(getattr(user, "team_id", 0)):
                warnings.append(f"无权访问文件: {file_id}")
                continue
            try:
                path = _ensure_outputs_path(self._settings.outputs_dir, file_id)
            except ValueError:
                warnings.append(f"文件路径非法: {file_id}")
                continue
            if not path.exists() or not path.is_file():
                warnings.append(f"文件不存在: {file_id}")
                continue
            resolved.append(
                {
                    "file_id": file_id,
                    "filename": _trim(item.get("filename")) or file_id,
                    "content_type": _trim(item.get("content_type")).lower(),
                    "size_bytes": int(item.get("size_bytes") or 0),
                    "path": path,
                }
            )
        return resolved, warnings
    async def _register_artifact(
        self,
        *,
        db,
        user,
        file_id: str,
        filename: str,
        content_type: str,
        size_bytes: int,
        kind: str,
        project_id: int | None,
    ) -> None:
        await db.execute(
            """
            INSERT OR IGNORE INTO file_records(
              file_id, team_id, user_id, project_id, session_id, kind, filename, content_type, size_bytes, created_at
            ) VALUES (?, ?, ?, ?, NULL, ?, ?, ?, ?, ?)
            """,
            (
                file_id,
                int(getattr(user, "team_id", 0)),
                int(getattr(user, "id", 0)),
                int(project_id) if project_id is not None and int(project_id) > 0 else None,
                kind,
                filename or file_id,
                content_type,
                int(size_bytes),
                utc_now_iso(),
            ),
        )
        await db.commit()
    async def _new_artifact(
        self,
        *,
        db,
        user,
        filename_hint: str,
        suffix: str,
        content: bytes,
        kind: str,
        detail: str,
        step: str,
        project_id: int | None,
    ) -> PipelineArtifact:
        safe_suffix = _safe_suffix(suffix if suffix.startswith(".") else f".{suffix}")
        file_id = f"{uuid4().hex}{safe_suffix or '.bin'}"
        path = _ensure_outputs_path(self._settings.outputs_dir, file_id)
        path.write_bytes(content)
        content_type = _content_type_by_suffix(path)
        await self._register_artifact(
            db=db,
            user=user,
            file_id=file_id,
            filename=_trim(filename_hint) or file_id,
            content_type=content_type,
            size_bytes=len(content),
            kind=kind,
            project_id=project_id,
        )
        token = create_download_token(settings=self._settings, file_id=file_id)
        return PipelineArtifact(
            file_id=file_id,
            filename=_trim(filename_hint) or file_id,
            kind=kind,
            detail=detail,
            step=step,
            download_url=abs_url(self._settings, f"/api/files/{file_id}?token={token}"),
            content_type=content_type,
            size_bytes=len(content),
        )
    async def _new_text_artifact(
        self,
        *,
        db,
        user,
        filename_hint: str,
        text: str,
        kind: str,
        detail: str,
        step: str,
        project_id: int | None,
    ) -> PipelineArtifact:
        suffix = _safe_suffix(filename_hint) or ".md"
        body = text if text.endswith("\n") else f"{text}\n"
        return await self._new_artifact(
            db=db,
            user=user,
            filename_hint=filename_hint,
            suffix=suffix,
            content=body.encode("utf-8"),
            kind=kind,
            detail=detail,
            step=step,
            project_id=project_id,
        )
    async def _artifact_from_meta(
        self,
        *,
        db,
        user,
        meta: dict[str, Any],
        kind: str,
        detail: str,
        step: str,
        project_id: int | None,
    ) -> PipelineArtifact | None:
        file_id = _trim(meta.get("file_id"))
        if not file_id:
            return None
        try:
            path = _ensure_outputs_path(self._settings.outputs_dir, file_id)
        except ValueError:
            return None
        if not path.exists() or not path.is_file():
            return None
        filename = _trim(meta.get("filename")) or file_id
        content_type = _trim(meta.get("content_type")) or _content_type_by_suffix(path)
        size_bytes = int(path.stat().st_size)
        await self._register_artifact(
            db=db,
            user=user,
            file_id=file_id,
            filename=filename,
            content_type=content_type,
            size_bytes=size_bytes,
            kind=kind,
            project_id=project_id,
        )
        token = create_download_token(settings=self._settings, file_id=file_id)
        return PipelineArtifact(
            file_id=file_id,
            filename=filename,
            kind=kind,
            detail=detail,
            step=step,
            download_url=abs_url(self._settings, f"/api/files/{file_id}?token={token}"),
            content_type=content_type,
            size_bytes=size_bytes,
        )
    def _which(self, name: str) -> str | None:
        value = shutil.which(name)
        return _trim(value) or None
    def _run_cmd(self, cmd: list[str], *, timeout: int = 180) -> tuple[bool, str]:
        try:
            proc = subprocess.run(cmd, capture_output=True, check=False, timeout=timeout)
        except Exception as e:
            return False, str(e)
        if proc.returncode == 0:
            return True, ""
        stderr = (proc.stderr or b"").decode("utf-8", errors="ignore").strip()
        stdout = (proc.stdout or b"").decode("utf-8", errors="ignore").strip()
        detail = stderr or stdout or f"exit={proc.returncode}"
        return False, detail[:600]
    async def _generate_text(self, *, system_prompt: str, user_prompt: str, fallback: str) -> str:
        api_key = _trim(self._settings.openai_api_key)
        if not api_key:
            return fallback
        try:
            provider = OpenAiProvider(
                api_key=api_key,
                base_url=self._settings.openai_base_url,
                verify_ssl=self._settings.openai_verify_ssl,
                outputs_dir=self._settings.outputs_dir,
            )
            res = await provider.complete(
                model=self._settings.model or "gpt-5.2",
                messages=[
                    ChatMessage(role="system", content=system_prompt),
                    ChatMessage(role="user", content=user_prompt),
                ],
                tools=[],
            )
            return _trim(res.assistant_text) or fallback
        except Exception:
            return fallback
    async def _describe_image_for_regeneration(
        self,
        *,
        prompt: str,
        image_inputs: list[dict[str, Any]],
        api_key: str,
        base_url: str,
        edit_scope: str,
        reference_strength: str,
        edit_region: dict[str, Any] | None,
    ) -> str:
        attachments: list[dict[str, Any]] = []
        for item in image_inputs[:4]:
            file_id = _trim(item.get("file_id"))
            if not file_id:
                continue
            attachments.append(
                {
                    "file_id": file_id,
                    "filename": _trim(item.get("filename")) or file_id,
                    "kind": "image",
                }
            )
        if not attachments:
            return prompt
        try:
            region_line = (
                f"局部编辑区域：x={edit_region.get('x')} y={edit_region.get('y')} width={edit_region.get('width')} height={edit_region.get('height')}\n"
                if isinstance(edit_region, dict) and edit_scope == "region"
                else ""
            )
            provider = OpenAiProvider(
                api_key=api_key,
                base_url=base_url,
                verify_ssl=self._settings.openai_verify_ssl,
                outputs_dir=self._settings.outputs_dir,
            )
            result = await provider.complete(
                model=_trim(self._settings.model) or "gpt-5.2",
                messages=[
                    ChatMessage(
                        role="system",
                        content=(
                            "你是一名图像编辑提示词工程师。"
                            "请根据用户的编辑要求和参考图，输出一段可直接用于文生图的完整提示词。"
                            "要求尽量保留主体、镜头视角、构图比例和关键视觉元素，并把用户要求的变化准确融入。"
                            "整体风格应接近原图的局部修饰或重绘，而不是完全换图。"
                            "只输出最终提示词，不要解释。"
                        ),
                    ),
                    ChatMessage(
                        role="user",
                        content=(
                            "请基于附图生成新的图片提示词。\n"
                            f"用户编辑要求：{prompt}\n"
                            f"编辑范围：{edit_scope}\n"
                            f"参考保真度：{reference_strength}\n"
                            f"{region_line}"
                            "输出要求：中文，适合直接提交给图片生成模型。"
                        ),
                        attachments=attachments,
                    ),
                ],
                tools=[],
            )
            described = _trim(result.assistant_text)
            return described or prompt
        except Exception:
            return prompt
    def _save_image_with_format(self, image: Image.Image, output_path: Path, fmt: str) -> tuple[bool, str]:
        try:
            if fmt == "jpg":
                image.convert("RGB").save(output_path, format="JPEG", quality=92, optimize=True)
            elif fmt == "webp":
                image.save(output_path, format="WEBP", quality=88, method=6)
            elif fmt == "png":
                image.save(output_path, format="PNG")
            elif fmt == "avif":
                ffmpeg = self._which("ffmpeg")
                if not ffmpeg:
                    return False, "未检测到 ffmpeg，无法输出 AVIF"
                tmp_png = output_path.with_suffix(".png")
                image.save(tmp_png, format="PNG")
                ok, err = self._run_cmd([ffmpeg, "-y", "-i", str(tmp_png), str(output_path)], timeout=120)
                try:
                    tmp_png.unlink(missing_ok=True)
                except Exception:
                    pass
                return ok and output_path.exists(), err
            else:
                image.save(output_path)
            return True, ""
        except Exception as e:
            return False, str(e)
    def _convert_file_for_office(self, *, path: Path, target_format: str) -> tuple[bytes | None, str | None, str]:
        fmt = _trim(target_format).lower()
        suffix = path.suffix.lower()
        if fmt not in {"txt", "md", "csv", "json", "xlsx", "png", "jpg", "webp", "mp3", "mp4"}:
            return None, None, f"不支持目标格式：{fmt}"
        try:
            if fmt in {"txt", "md"}:
                if suffix == ".docx":
                    text = _text_from_docx(path)
                elif suffix == ".xlsx":
                    text = _xlsx_to_csv_text(path)
                else:
                    text = path.read_text(encoding="utf-8", errors="ignore")
                return text.encode("utf-8"), f".{fmt}", ""
            if fmt == "csv":
                if suffix == ".xlsx":
                    text = _xlsx_to_csv_text(path)
                elif suffix == ".json":
                    text = _json_text_to_csv(path.read_text(encoding="utf-8", errors="ignore"))
                else:
                    text = path.read_text(encoding="utf-8", errors="ignore")
                return text.encode("utf-8"), ".csv", ""
            if fmt == "json":
                if suffix == ".csv":
                    reader = csv.DictReader(io.StringIO(path.read_text(encoding="utf-8", errors="ignore")))
                    payload = list(reader)
                else:
                    payload = {"text": path.read_text(encoding="utf-8", errors="ignore")}
                return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"), ".json", ""
            if fmt == "xlsx":
                source = path.read_text(encoding="utf-8", errors="ignore") if suffix != ".xlsx" else _xlsx_to_csv_text(path)
                return _csv_text_to_xlsx_bytes(source), ".xlsx", ""
            if fmt in {"png", "jpg", "webp"} and suffix in _IMAGE_EXTS:
                image = Image.open(path)
                buf = io.BytesIO()
                save_format = "JPEG" if fmt == "jpg" else fmt.upper()
                if fmt == "jpg":
                    image.convert("RGB").save(buf, format=save_format, quality=92)
                else:
                    image.save(buf, format=save_format)
                return buf.getvalue(), f".{fmt}", ""
            if fmt in {"mp3", "mp4"} and (suffix in _AUDIO_EXTS or suffix in _VIDEO_EXTS):
                ffmpeg = self._which("ffmpeg")
                if not ffmpeg:
                    return None, None, "未检测到 ffmpeg"
                tmp_name = f"office-convert-{uuid4().hex[:8]}.{fmt}"
                tmp_path = _ensure_outputs_path(self._settings.outputs_dir, tmp_name)
                if fmt == "mp3":
                    ok, err = self._run_cmd(
                        [ffmpeg, "-y", "-i", str(path), "-vn", "-c:a", "libmp3lame", "-b:a", "128k", str(tmp_path)],
                        timeout=180,
                    )
                else:
                    ok, err = self._run_cmd(
                        [ffmpeg, "-y", "-i", str(path), "-c:v", "libx264", "-preset", "veryfast", "-crf", "24", "-c:a", "aac", str(tmp_path)],
                        timeout=240,
                    )
                if not ok or not tmp_path.exists():
                    return None, None, err or "转换失败"
                payload = tmp_path.read_bytes()
                try:
                    tmp_path.unlink(missing_ok=True)
                except Exception:
                    pass
                return payload, f".{fmt}", ""
        except Exception as e:
            return None, None, str(e)
        return None, None, f"暂不支持从 {suffix or 'unknown'} 转到 {fmt}"
    async def _create_ai_image_artifact(
        self,
        *,
        db,
        user,
        prompt: str,
        operation: str,
        image_inputs: list[dict[str, Any]],
        size: str,
        quality: str,
        background: str,
        output_format: str,
        input_fidelity: str,
        edit_scope: str,
        reference_strength: str,
        edit_region: dict[str, Any] | None,
        project_id: int | None,
    ) -> tuple[PipelineArtifact | None, dict[str, Any] | None, list[str]]:
        warnings: list[str] = []
        prompt_text = _trim(prompt)
        if not prompt_text:
            warnings.append("未填写 prompt，AI 图片生成/编辑已跳过。")
            return None, None, warnings
        preferred_provider = _trim(getattr(self._settings, "provider", "")).lower()
        openai_api_key = _trim(self._settings.openai_api_key)
        glm_api_key = _trim(getattr(self._settings, "glm_api_key", None))
        use_glm_generate = operation == "generate" and bool(glm_api_key)
        if operation != "generate" and preferred_provider == "glm":
            if openai_api_key:
                warnings.append("GLM 图片编辑暂未接入，已自动回退到 OpenAI 图片编辑链路。")
            elif glm_api_key:
                warnings.append("GLM 当前仅接入免费文生图；图片编辑请配置 OPENAI_API_KEY 后重试。")
                return None, None, warnings
        provider_name = "glm" if use_glm_generate else "openai"
        if provider_name == "glm":
            api_key = glm_api_key
            model = _trim(getattr(self._settings, "glm_image_model", "")) or "cogview-3-flash"
            base_url = _normalize_openai_base_url(self._settings.glm_base_url)
            endpoint = f"{base_url}/images/generations"
        else:
            api_key = openai_api_key
            if not api_key:
                if glm_api_key:
                    warnings.append("GLM 当前仅接入免费文生图；图片编辑请改用生成模式，或配置 OPENAI_API_KEY。")
                else:
                    warnings.append("未配置 OPENAI_API_KEY 或 GLM_API_KEY，AI 图片生成/编辑已跳过。")
                return None, None, warnings
            model = _trim(self._settings.openai_image_model) or "gpt-image-1"
            base_url = _normalize_openai_base_url(self._settings.openai_base_url)
            endpoint = f"{base_url}/images/generations" if operation == "generate" else f"{base_url}/images/edits"
        timeout = httpx.Timeout(180.0)
        bytes_payload: bytes | None = None
        remote_url: str | None = None
        revised_prompt: str | None = None
        content_type: str | None = None
        postprocess_source_path: Path | None = None
        async with httpx.AsyncClient(timeout=timeout) as client:
            async def _post_image_request(url: str, *, headers: dict[str, str], data=None, json=None, files=None):  # noqa: ANN001
                last_response = None
                for attempt in range(3):
                    response = await client.post(url, headers=headers, data=data, json=json, files=files)
                    last_response = response
                    if not _is_retryable_image_error(response) or attempt == 2:
                        return response
                    await asyncio.sleep(1.2 * (attempt + 1))
                return last_response

            if provider_name == "glm":
                payload = {
                    "model": model,
                    "prompt": prompt_text,
                    "size": size,
                    "quality": _glm_image_quality(quality),
                }
                response = await _post_image_request(
                    endpoint,
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json",
                    },
                    json=payload,
                )
            elif operation == "generate":
                payload = {
                    "model": model,
                    "prompt": prompt_text,
                    "size": size,
                    "quality": quality,
                    "background": background,
                    "n": 1,
                }
                if output_format != "png":
                    payload["output_format"] = output_format
                response = await _post_image_request(
                    endpoint,
                    headers={
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json",
                    },
                    json=payload,
                )
            else:
                usable_inputs = [
                    item for item in image_inputs
                    if isinstance(item.get("path"), Path) and item["path"].exists() and item["path"].is_file()
                ][:4]
                usable_paths = [item["path"] for item in usable_inputs]
                if not usable_paths:
                    warnings.append("编辑图片至少需要上传一张输入图片。")
                    return None, None, warnings
                data = {
                    "model": model,
                    "prompt": prompt_text,
                    "size": size,
                    "quality": quality,
                    "background": background,
                    "input_fidelity": input_fidelity,
                }
                if output_format != "png":
                    data["output_format"] = output_format
                files: list[tuple[str, tuple[str, bytes, str]]] = []
                for image_path in usable_paths:
                    ctype = _content_type_by_suffix(image_path)
                    field_name = "image[]" if len(usable_paths) > 1 else "image"
                    files.append((field_name, (image_path.name, image_path.read_bytes(), ctype)))
                response = await _post_image_request(
                    endpoint,
                    headers={"Authorization": f"Bearer {api_key}"},
                    data=data,
                    files=files,
                )
                if response.status_code >= 400 and _supports_image_edit_fallback(response):
                    fallback_prompt = await self._describe_image_for_regeneration(
                        prompt=prompt_text,
                        image_inputs=usable_inputs,
                        api_key=api_key,
                        base_url=base_url,
                        edit_scope=edit_scope,
                        reference_strength=reference_strength,
                        edit_region=edit_region,
                    )
                    primary_source = usable_paths[0]
                    fallback_size = _pick_closest_image_size(primary_source.width if hasattr(primary_source, "width") else 0, 0)
                    try:
                        with Image.open(primary_source) as source_img:
                            fallback_size = _pick_closest_image_size(source_img.width, source_img.height)
                    except Exception:
                        fallback_size = size
                    fallback_api_key = glm_api_key or api_key
                    fallback_provider = "glm" if glm_api_key else "openai"
                    if fallback_provider == "glm":
                        fallback_model = _trim(getattr(self._settings, "glm_image_model", "")) or "cogview-3-flash"
                        fallback_base_url = _normalize_openai_base_url(self._settings.glm_base_url)
                        fallback_payload = {
                            "model": fallback_model,
                            "prompt": fallback_prompt,
                            "size": fallback_size,
                            "quality": _glm_image_quality(quality),
                        }
                    else:
                        fallback_model = model
                        fallback_base_url = base_url
                        fallback_payload = {
                            "model": fallback_model,
                            "prompt": fallback_prompt,
                            "size": fallback_size,
                            "quality": quality,
                            "background": background,
                            "n": 1,
                            **({"output_format": output_format} if output_format != "png" else {}),
                        }
                    warnings.append(f"图片编辑接口不可用，已自动回退为参考图理解后再生成（{fallback_provider.upper()}）。")
                    postprocess_source_path = primary_source
                    response = await _post_image_request(
                        f"{fallback_base_url}/images/generations",
                        headers={
                            "Authorization": f"Bearer {fallback_api_key}",
                            "Content-Type": "application/json",
                        },
                        json=fallback_payload,
                    )
                    if fallback_prompt != prompt_text:
                        revised_prompt = fallback_prompt
            if response.status_code >= 400:
                warnings.append(f"AI 图片接口失败：{_extract_http_error_text(response)}")
                return None, None, warnings
            payload = response.json()
            bytes_payload, remote_url, revised_prompt = _first_image_payload(payload if isinstance(payload, dict) else {})
            if remote_url and not bytes_payload:
                remote_resp = await client.get(remote_url)
                if remote_resp.status_code >= 400:
                    warnings.append(f"拉取生成图片失败：{_extract_http_error_text(remote_resp)}")
                    return None, None, warnings
                bytes_payload = remote_resp.content
                content_type = _trim(remote_resp.headers.get("content-type")) or None
        if not bytes_payload:
            warnings.append("AI 图片接口未返回可用图片数据。")
            return None, None, warnings
        if operation == "edit" and postprocess_source_path is not None:
            fitted = _fit_generated_image_to_source(bytes_payload, postprocess_source_path)
            if fitted is not None:
                bytes_payload, content_type = fitted
            scoped = _composite_generated_edit(
                source_path=postprocess_source_path,
                generated_bytes=bytes_payload,
                edit_scope=edit_scope,
                reference_strength=reference_strength,
                edit_region=edit_region,
            )
            if scoped is not None:
                bytes_payload, content_type = scoped
        inferred_content_type = _infer_image_content_type_from_bytes(bytes_payload)
        if inferred_content_type:
            content_type = inferred_content_type
        suffix, default_type = _image_suffix_and_content_type_by_header(content_type, output_format)
        filename_hint = f"ai-image-{operation}-{uuid4().hex[:8]}{suffix}"
        detail = (
            "GLM 图片生成"
            if provider_name == "glm" and operation == "generate"
            else ("AI 图片生成" if operation == "generate" else "智能重绘/重生成")
        )
        artifact = await self._new_artifact(
            db=db,
            user=user,
            filename_hint=filename_hint,
            suffix=suffix,
            content=bytes_payload,
            kind="vision",
            detail=detail,
            step="image_generate" if operation == "generate" else "image_edit",
            project_id=project_id,
        )
        meta = {
            "file_id": artifact.file_id,
            "filename": artifact.filename,
            "content_type": content_type or default_type,
            "size_bytes": artifact.size_bytes,
            "path": _ensure_outputs_path(self._settings.outputs_dir, artifact.file_id),
            "provider": provider_name,
            "model": model,
        }
        if revised_prompt and revised_prompt != prompt_text:
            warnings.append(f"模型修订提示词：{revised_prompt}")
        return artifact, meta, warnings
    async def _run_vision_pipeline(
        self,
        *,
        payload: dict[str, Any],
        mode: str,
        user,
        db,
        project_id: int | None,
    ) -> tuple[str, list[PipelineArtifact], list[str]]:
        file_ids = [str(v) for v in (payload.get("input_file_ids") or []) if str(v).strip()]
        inputs, warnings = await self._resolve_input_files(file_ids=file_ids, user=user, db=db)
        artifacts: list[PipelineArtifact] = []
        formats = _normalize_target_formats(payload.get("target_formats"), defaults=["webp", "jpg", "png"])
        resize = payload.get("resize") if isinstance(payload.get("resize"), dict) else {}
        width = _safe_int(resize.get("width") if resize else 0, 0)
        enhance = bool(payload.get("enhance", True))
        prompt = _trim(payload.get("prompt"))
        operation = _normalize_vision_operation(payload.get("operation"), has_inputs=bool(inputs))
        image_size = _normalize_image_size(payload.get("size"))
        image_quality = _normalize_image_quality(payload.get("quality"))
        image_background = _normalize_image_background(payload.get("background"))
        image_output = _normalize_image_output(payload.get("output_format"))
        input_fidelity = _normalize_input_fidelity(payload.get("input_fidelity"))
        edit_scope = _normalize_edit_scope(payload.get("edit_scope"))
        reference_strength = _normalize_reference_strength(payload.get("reference_strength"))
        edit_region = payload.get("edit_region") if isinstance(payload.get("edit_region"), dict) else None
        crop_presets = _normalize_crop_presets(payload.get("crop_presets"))
        watermark_text = _trim(payload.get("watermark_text"))
        watermark_position = _normalize_watermark_position(payload.get("watermark_position"))
        filter_preset = _trim(payload.get("filter_preset")).lower() or "business"
        transparent_background = bool(payload.get("transparent_background", False))
        processing_inputs = inputs
        ai_artifact: PipelineArtifact | None = None
        ai_meta: dict[str, Any] | None = None
        if prompt:
            ai_artifact, ai_meta, ai_warnings = await self._create_ai_image_artifact(
                db=db,
                user=user,
                prompt=prompt,
                operation=operation,
                image_inputs=inputs,
                size=image_size,
                quality=image_quality,
                background=image_background,
                output_format=image_output,
                input_fidelity=input_fidelity,
                edit_scope=edit_scope,
                reference_strength=reference_strength,
                edit_region=edit_region,
                project_id=project_id,
            )
            warnings.extend(ai_warnings)
            if ai_artifact and ai_meta:
                artifacts.append(ai_artifact)
                processing_inputs = [ai_meta]
        elif not inputs:
            warnings.append("未提供 prompt 或输入图片，已仅生成执行建议文档。")
        should_post_process = _should_post_process(enhance=enhance, width=width, formats=formats, primary_format="jpg" if image_output == "jpeg" else image_output)
        if processing_inputs and should_post_process:
            for item in processing_inputs:
                path: Path = item["path"]
                suffix = path.suffix.lower()
                ctype = str(item.get("content_type") or "")
                if suffix not in _IMAGE_EXTS and not ctype.startswith("image/"):
                    continue
                try:
                    base_image = Image.open(path)
                except Exception as e:
                    warnings.append(f"图片读取失败：{item['filename']} ({e})")
                    continue
                processed = _resize_image(base_image, width=width)
                processed = _apply_filter_preset(processed, filter_preset, enhance=enhance)
                if transparent_background or image_background == "transparent":
                    processed = _apply_transparent_background(processed)
                if watermark_text:
                    processed = _apply_watermark(processed, text=watermark_text, position=watermark_position)
                variants: list[tuple[str, Image.Image, str]] = [("base", processed, "图像增强与滤镜")]
                for preset in crop_presets:
                    variants.append((preset.replace(":", "x"), _crop_image_to_ratio(processed, preset), f"裁剪 {preset}"))
                base = f"{_safe_stem(path.stem)}-{mode}"
                for variant_key, variant_image, variant_detail in variants:
                    for fmt in formats:
                        output_name = f"{base}-{variant_key}-{uuid4().hex[:8]}.{fmt}"
                        output_path = _ensure_outputs_path(self._settings.outputs_dir, output_name)
                        ok, err = self._save_image_with_format(variant_image, output_path, fmt)
                        if not ok or not output_path.exists():
                            warnings.append(f"图片处理失败：{item['filename']} -> {fmt} ({err})")
                            continue
                        artifact = await self._artifact_from_meta(
                            db=db,
                            user=user,
                            meta={"file_id": output_name, "filename": output_name, "content_type": _content_type_by_suffix(output_path)},
                            kind="vision",
                            detail=f"{variant_detail} / {fmt.upper()}",
                            step="image_process" if variant_key == "base" else "image_crop",
                            project_id=project_id,
                        )
                        if artifact:
                            artifacts.append(artifact)
        if not artifacts and processing_inputs and not ai_artifact:
            warnings.append("当前配置未产出新图片，请检查格式/增强参数。")
        report_lines = [
            "# 视觉设计与图像处理 Pipeline 报告",
            "",
            f"- 部署模式：{mode}",
            f"- 执行动作：{'AI 图片生成' if operation == 'generate' else 'AI 图片编辑'}" if prompt else "- 执行动作：图片处理",
            f"- Prompt：{prompt}" if prompt else "- Prompt：未填写",
            f"- 输入文件数：{len(inputs)}",
            f"- 目标格式：{', '.join(formats)}",
            f"- 裁剪预设：{', '.join(crop_presets) if crop_presets else '未启用'}",
            f"- 编辑范围：{edit_scope}",
            f"- 参考保真度：{reference_strength}",
            f"- 滤镜预设：{filter_preset}",
            f"- 水印：{watermark_text or '未启用'}",
            f"- 透明背景：{'开启' if (transparent_background or image_background == 'transparent') else '关闭'}",
            f"- 输出产物数：{len(artifacts)}",
            "",
            "## 执行内容",
        ]
        if prompt:
            report_lines.append("- AI 出图：已调用图像模型执行生成/编辑。")
        report_lines += [
            "- 图片编辑/裁剪：支持按预设比例生成裁剪版本。",
            "- 滤镜与增强：执行对比度、饱和度、锐化与商务/暖色/戏剧化预设。",
            "- 水印与透明背景：支持文字水印和近白底透明化导出。",
            "- 格式转换：按目标格式输出，便于云端分发与边缘落地。",
        ]
        if warnings:
            report_lines += ["", "## 告警", *[f"- {w}" for w in warnings]]
        report = "\n".join(report_lines)
        artifacts.insert(
            0,
            await self._new_text_artifact(
                db=db,
                user=user,
                filename_hint=f"vision-pipeline-report-{uuid4().hex[:8]}.md",
                text=report,
                kind="vision",
                detail="视觉 pipeline 执行报告",
                step="summary",
                project_id=project_id,
            ),
        )
        return report, artifacts, warnings
    async def _run_media_pipeline(
        self,
        *,
        payload: dict[str, Any],
        mode: str,
        user,
        db,
        project_id: int | None,
    ) -> tuple[str, list[PipelineArtifact], list[str]]:
        file_ids = [str(v) for v in (payload.get("input_file_ids") or []) if str(v).strip()]
        merge_file_ids = [str(v) for v in (payload.get("merge_input_file_ids") or []) if str(v).strip()]
        inputs, warnings = await self._resolve_input_files(file_ids=file_ids, user=user, db=db)
        if merge_file_ids:
            extra_inputs, extra_warnings = await self._resolve_input_files(file_ids=merge_file_ids, user=user, db=db)
            inputs.extend(extra_inputs)
            warnings.extend(extra_warnings)
        artifacts: list[PipelineArtifact] = []
        ffmpeg = self._which("ffmpeg")
        clip_seconds = min(max(5, int(payload.get("clip_seconds") or 15)), 180)
        output_resolutions = [str(v).strip() for v in (payload.get("output_resolutions") or ["1080", "720"]) if str(v).strip()]
        output_resolutions = [v for v in output_resolutions if v in {"1080", "720", "480"}] or ["1080", "720"]
        target_containers = [str(v).strip().lower() for v in (payload.get("target_containers") or ["mp4"]) if str(v).strip()]
        target_containers = [v for v in target_containers if v in {"mp4", "hls"}] or ["mp4"]
        video_codec = _trim(payload.get("video_codec")).lower() or "h264"
        if video_codec not in {"h264", "h265"}:
            video_codec = "h264"
        subtitle_text = _trim(payload.get("subtitle_text"))
        cover_at_second = max(0, _safe_int(payload.get("cover_at_second"), 1))
        trim_silence = bool(payload.get("trim_silence", True))
        if not inputs:
            warnings.append("未提供可处理音视频，已仅生成执行建议文档。")
        if not ffmpeg:
            warnings.append("未检测到 ffmpeg，音视频处理步骤已跳过。")
        if ffmpeg and inputs:
            video_inputs: list[Path] = []
            for item in inputs:
                path: Path = item["path"]
                suffix = path.suffix.lower()
                ctype = str(item.get("content_type") or "")
                is_video = suffix in _VIDEO_EXTS or ctype.startswith("video/")
                is_audio = suffix in _AUDIO_EXTS or ctype.startswith("audio/")
                base = f"{_safe_stem(path.stem, max_len=36)}-{mode}-{uuid4().hex[:6]}"
                vcodec = "libx265" if video_codec == "h265" else "libx264"
                video_crf = "28" if mode == "edge" else "24"
                audio_bitrate = "96k" if mode == "edge" else "128k"
                if is_video:
                    video_inputs.append(path)
                    for resolution in output_resolutions:
                        height = {"1080": 1080, "720": 720, "480": 480}[resolution]
                        vf_parts = [f"scale=-2:{height}:flags=lanczos"]
                        if subtitle_text:
                            escaped = subtitle_text.replace("\\", "\\\\").replace(":", "\\:").replace("'", "\\'")
                            vf_parts.append(f"drawtext=text='{escaped}':fontcolor=white:fontsize=24:x=(w-text_w)/2:y=h-(text_h*2)")
                        for container in target_containers:
                            if container == "mp4":
                                mp4_name = f"{base}-{resolution}-{video_codec}.mp4"
                                mp4_path = _ensure_outputs_path(self._settings.outputs_dir, mp4_name)
                                ok, err = self._run_cmd(
                                    [
                                        ffmpeg,
                                        "-y",
                                        "-i",
                                        str(path),
                                        "-vf",
                                        ",".join(vf_parts),
                                        "-c:v",
                                        vcodec,
                                        "-preset",
                                        "veryfast",
                                        "-crf",
                                        video_crf,
                                        "-c:a",
                                        "aac",
                                        "-b:a",
                                        audio_bitrate,
                                        str(mp4_path),
                                    ],
                                    timeout=300,
                                )
                                if ok and mp4_path.exists():
                                    artifact = await self._artifact_from_meta(
                                        db=db,
                                        user=user,
                                        meta={"file_id": mp4_name, "filename": mp4_name, "content_type": "video/mp4"},
                                        kind="media",
                                        detail=f"视频转码 {resolution}p {video_codec.upper()}",
                                        step="video_transcode",
                                        project_id=project_id,
                                    )
                                    if artifact:
                                        artifacts.append(artifact)
                                else:
                                    warnings.append(f"视频转码失败：{item['filename']} {resolution}p ({err})")
                            else:
                                hls_dir = _ensure_outputs_path(self._settings.outputs_dir, f"{base}-{resolution}-hls-{uuid4().hex[:6]}")
                                hls_dir.mkdir(parents=True, exist_ok=True)
                                playlist = hls_dir / "index.m3u8"
                                ok, err = self._run_cmd(
                                    [
                                        ffmpeg,
                                        "-y",
                                        "-i",
                                        str(path),
                                        "-vf",
                                        ",".join(vf_parts),
                                        "-c:v",
                                        vcodec,
                                        "-preset",
                                        "veryfast",
                                        "-crf",
                                        video_crf,
                                        "-c:a",
                                        "aac",
                                        "-b:a",
                                        audio_bitrate,
                                        "-f",
                                        "hls",
                                        "-hls_time",
                                        "4",
                                        "-hls_playlist_type",
                                        "vod",
                                        str(playlist),
                                    ],
                                    timeout=360,
                                )
                                if ok and playlist.exists():
                                    zip_name = f"{base}-{resolution}-hls.zip"
                                    shutil.make_archive(str(_ensure_outputs_path(self._settings.outputs_dir, zip_name).with_suffix("")), "zip", root_dir=hls_dir)
                                    artifact = await self._artifact_from_meta(
                                        db=db,
                                        user=user,
                                        meta={"file_id": zip_name, "filename": zip_name, "content_type": "application/zip"},
                                        kind="media",
                                        detail=f"HLS 分发包 {resolution}p {video_codec.upper()}",
                                        step="video_hls",
                                        project_id=project_id,
                                    )
                                    if artifact:
                                        artifacts.append(artifact)
                                else:
                                    warnings.append(f"HLS 输出失败：{item['filename']} {resolution}p ({err})")
                    clip_name = f"{base}-clip.mp4"
                    clip_path = _ensure_outputs_path(self._settings.outputs_dir, clip_name)
                    ok, err = self._run_cmd(
                        [
                            ffmpeg,
                            "-y",
                            "-ss",
                            "0",
                            "-i",
                            str(path),
                            "-t",
                            str(clip_seconds),
                            "-c:v",
                            vcodec,
                            "-preset",
                            "veryfast",
                            "-crf",
                            video_crf,
                            "-c:a",
                            "aac",
                            "-b:a",
                            audio_bitrate,
                            str(clip_path),
                        ],
                        timeout=180,
                    )
                    if ok and clip_path.exists():
                        artifact = await self._artifact_from_meta(
                            db=db,
                            user=user,
                            meta={"file_id": clip_name, "filename": clip_name, "content_type": "video/mp4"},
                            kind="media",
                            detail=f"视频剪辑 {clip_seconds}s",
                            step="video_clip",
                            project_id=project_id,
                        )
                        if artifact:
                            artifacts.append(artifact)
                    else:
                        warnings.append(f"视频剪辑失败：{item['filename']} ({err})")
                    cover_name = f"{base}-cover.jpg"
                    cover_path = _ensure_outputs_path(self._settings.outputs_dir, cover_name)
                    ok, err = self._run_cmd([ffmpeg, "-y", "-ss", str(cover_at_second), "-i", str(path), "-frames:v", "1", str(cover_path)], timeout=120)
                    if ok and cover_path.exists():
                        artifact = await self._artifact_from_meta(
                            db=db,
                            user=user,
                            meta={"file_id": cover_name, "filename": cover_name, "content_type": "image/jpeg"},
                            kind="media",
                            detail=f"封面抽帧 {cover_at_second}s",
                            step="video_cover",
                            project_id=project_id,
                        )
                        if artifact:
                            artifacts.append(artifact)
                    else:
                        warnings.append(f"封面抽帧失败：{item['filename']} ({err})")
                if is_video or is_audio:
                    mp3_name = f"{base}-denoise.mp3"
                    mp3_path = _ensure_outputs_path(self._settings.outputs_dir, mp3_name)
                    audio_filter = "afftdn,loudnorm=I=-16:TP=-1.5:LRA=11"
                    if trim_silence:
                        audio_filter = f"silenceremove=start_periods=1:start_duration=0.2:start_threshold=-45dB,{audio_filter}"
                    ok, err = self._run_cmd(
                        [
                            ffmpeg,
                            "-y",
                            "-i",
                            str(path),
                            "-vn",
                            "-af",
                            audio_filter,
                            "-c:a",
                            "libmp3lame",
                            "-b:a",
                            "128k",
                            str(mp3_path),
                        ],
                        timeout=180,
                    )
                    if ok and mp3_path.exists():
                        artifact = await self._artifact_from_meta(
                            db=db,
                            user=user,
                            meta={"file_id": mp3_name, "filename": mp3_name, "content_type": "audio/mpeg"},
                            kind="media",
                            detail="音频降噪与标准化",
                            step="audio_denoise",
                            project_id=project_id,
                        )
                        if artifact:
                            artifacts.append(artifact)
                    else:
                        warnings.append(f"音频降噪失败：{item['filename']} ({err})")
            if len(video_inputs) >= 2:
                concat_name = f"merged-{mode}-{uuid4().hex[:8]}.mp4"
                concat_path = _ensure_outputs_path(self._settings.outputs_dir, concat_name)
                concat_list_name = f"concat-{uuid4().hex[:8]}.txt"
                concat_list_path = _ensure_outputs_path(self._settings.outputs_dir, concat_list_name)
                concat_list_path.write_text("".join(f"file '{p.as_posix()}'\n" for p in video_inputs), encoding="utf-8")
                ok, err = self._run_cmd([ffmpeg, "-y", "-f", "concat", "-safe", "0", "-i", str(concat_list_path), "-c", "copy", str(concat_path)], timeout=300)
                try:
                    concat_list_path.unlink(missing_ok=True)
                except Exception:
                    pass
                if ok and concat_path.exists():
                    artifact = await self._artifact_from_meta(
                        db=db,
                        user=user,
                        meta={"file_id": concat_name, "filename": concat_name, "content_type": "video/mp4"},
                        kind="media",
                        detail="视频合并输出",
                        step="video_merge",
                        project_id=project_id,
                    )
                    if artifact:
                        artifacts.append(artifact)
                else:
                    warnings.append(f"视频合并失败：{err}")
        report_lines = [
            "# 视频与音频处理 Pipeline 报告",
            "",
            f"- 部署模式：{mode}",
            f"- 输入文件数：{len(inputs)}",
            f"- 剪辑时长：{clip_seconds}s",
            f"- 输出分辨率：{', '.join(output_resolutions)}",
            f"- 输出容器：{', '.join(target_containers)}",
            f"- 视频编码：{video_codec.upper()}",
            f"- 字幕烧录：{'开启' if subtitle_text else '关闭'}",
            f"- 静音裁剪：{'开启' if trim_silence else '关闭'}",
            f"- 输出产物数：{len(artifacts)}",
            "",
            "## 执行内容",
            "- 视频剪辑：输出可快速预览片段。",
            "- 视频转码：支持多分辨率、H.264/H.265 与 MP4/HLS。",
            "- 视频合并与封面：支持多片段合并与关键帧封面抽取。",
            "- 字幕与音频：支持字幕烧录、降噪、静音裁剪与响度标准化。",
        ]
        if warnings:
            report_lines += ["", "## 告警", *[f"- {w}" for w in warnings]]
        report = "\n".join(report_lines)
        artifacts.insert(
            0,
            await self._new_text_artifact(
                db=db,
                user=user,
                filename_hint=f"media-pipeline-report-{uuid4().hex[:8]}.md",
                text=report,
                kind="media",
                detail="音视频 pipeline 执行报告",
                step="summary",
                project_id=project_id,
            ),
        )
        return report, artifacts, warnings
    async def _run_content_pipeline(
        self,
        *,
        payload: dict[str, Any],
        mode: str,
        user,
        db,
        project_id: int | None,
    ) -> tuple[str, list[PipelineArtifact], list[str]]:
        topic = _trim(payload.get("topic")) or "AI+通用工具集发布方案"
        audience = _trim(payload.get("audience")) or "小团队与一人公司（OPC）"
        tone = _trim(payload.get("tone")) or "专业、清晰、可落地"
        points_raw = payload.get("key_points")
        points = [str(v).strip() for v in (points_raw or []) if str(v).strip()]
        if not points:
            points = ["云边协同架构", "多模态处理能力", "可切换技能化 Pipeline", "可观测与可回滚交付"]
        variants = min(max(1, _safe_int(payload.get("variants"), 3)), 6)
        languages = [str(v).strip() for v in (payload.get("languages") or ["zh-CN"]) if str(v).strip()] or ["zh-CN"]
        rewrite_styles = [str(v).strip() for v in (payload.get("rewrite_styles") or ["专业版", "亲和版", "冲击版"]) if str(v).strip()]
        brand_terms = [str(v).strip() for v in (payload.get("brand_terms") or []) if str(v).strip()]
        source_text = _trim(payload.get("source_text"))
        review_focus = [str(v).strip() for v in (payload.get("review_focus") or []) if str(v).strip()]
        warnings: list[str] = []
        artifacts: list[PipelineArtifact] = []
        variants_buf: list[str] = []
        style_names = rewrite_styles or [f"版本 {idx}" for idx in range(1, variants + 1)]
        for idx in range(variants):
            style_name = style_names[idx % len(style_names)]
            base_copy = "\n".join(
                [
                    f"### {style_name}",
                    f"- 标题：{topic}",
                    f"- 受众：{audience}",
                    f"- 语气：{tone}",
                    f"- 品牌词：{'、'.join(brand_terms) if brand_terms else '无'}",
                    f"- 卖点：{points[idx % len(points)]}",
                    f"- 文案：{topic} 通过{points[idx % len(points)]}，把视觉、音视频、内容和办公交付收敛成统一流程。",
                ]
            )
            variants_buf.append(base_copy)
        revised_text = source_text
        if source_text:
            fallback_rewrite = "\n".join(
                [
                    f"# {topic} 文档修订稿",
                    "",
                    "## 摘要",
                    f"{topic} 面向 {audience}，当前重点是把 {', '.join(points[:3])} 组织成可执行方案。",
                    "",
                    "## 修订正文",
                    source_text,
                    "",
                    "## 收尾建议",
                    "补充验收标准、时间表与责任分工，形成正式交付版本。",
                ]
            )
            revised_text = await self._generate_text(
                system_prompt="你是一名企业内容编辑，请把输入内容整理成更清晰的中文方案文稿，保持事实，不夸张。",
                user_prompt=f"主题：{topic}\n受众：{audience}\n语气：{tone}\n品牌词：{', '.join(brand_terms)}\n请整理以下文稿：\n{source_text}",
                fallback=fallback_rewrite,
            )
        localized_sections: list[str] = []
        for lang in languages:
            localized_text = revised_text or "\n\n".join(variants_buf)
            if lang.lower() not in {"zh", "zh-cn", "zh-hans"}:
                localized_text = await self._generate_text(
                    system_prompt="You are a localization writer. Translate the content while keeping markdown structure.",
                    user_prompt=f"Target language: {lang}\nContent:\n{localized_text}",
                    fallback=_offline_translate(localized_text, lang),
                )
            localized_sections.append(f"## {_language_tag(lang)}\n\n{localized_text}")
        suggestions = _review_suggestions(revised_text or "\n".join(variants_buf), review_focus, brand_terms)
        summary_lines = _version_summary(source_text, revised_text) if source_text else ["本次以新建内容为主，已生成多版本文案与交付文档。"]
        markdown = "\n".join(
            [
                f"# {topic}",
                "",
                f"- 目标受众：{audience}",
                f"- 文案语气：{tone}",
                f"- 部署模式：{mode}",
                f"- 语言：{', '.join(languages)}",
                f"- 文案版本数：{variants}",
                "",
                "## 核心卖点",
                *[f"- {item}" for item in points],
                "",
                "## 文案版本",
                *variants_buf,
                "",
                "## 多语输出",
                *localized_sections,
                "",
                "## 审阅建议",
                *[f"- {item}" for item in suggestions],
                "",
                "## 版本摘要",
                *[f"- {item}" for item in summary_lines],
            ]
        )
        artifacts.append(
            await self._new_text_artifact(
                db=db,
                user=user,
                filename_hint=f"content-copy-pack-{uuid4().hex[:8]}.md",
                text=markdown,
                kind="content",
                detail="内容创作文案包",
                step="copywriting",
                project_id=project_id,
            )
        )
        plan_docx = _build_docx_bytes(
            f"{topic} 执行方案",
            [
                ("项目目标", f"目标受众：{audience}\n语气：{tone}\n部署模式：{mode}"),
                ("能力范围", "\n".join(f"- {item}" for item in points)),
                ("修订与审阅", "\n".join(suggestions)),
                ("版本摘要", "\n".join(summary_lines)),
            ],
        )
        artifacts.append(
            await self._new_artifact(
                db=db,
                user=user,
                filename_hint=f"content-plan-{uuid4().hex[:8]}.docx",
                suffix=".docx",
                content=plan_docx,
                kind="content",
                detail="方案文档（DOCX）",
                step="document",
                project_id=project_id,
            )
        )
        slides = [
            {"title": "项目目标", "bullets": [topic, audience, "构建可切换可执行的技能化流程"]},
            {"title": "能力地图", "bullets": points[:4]},
            {"title": "执行与验收", "bullets": ["输入输出定义", "云边资源分配", "回滚与审计策略"]},
            {"title": "审阅与版本", "bullets": suggestions[:3] or ["已生成审阅建议与版本摘要"]},
        ]
        try:
            ppt_meta = await self._doc_service.create_pptx(
                title=f"{topic} 汇报",
                slides=slides,
                style="template_jetlinks",
                layout_mode="auto",
            )
            artifact = await self._artifact_from_meta(
                db=db,
                user=user,
                meta=ppt_meta,
                kind="content",
                detail="内容提案 PPT",
                step="presentation",
                project_id=project_id,
            )
            if artifact:
                artifacts.append(artifact)
        except Exception as e:
            warnings.append(f"PPT 生成失败：{e}")
        quote_items = []
        for idx, point in enumerate(points[:4], 1):
            quote_items.append(
                {
                    "name": point or f"模块 {idx}",
                    "quantity": 1,
                    "unit_price": float(3000 + idx * 1200),
                    "unit": "项",
                    "note": "含 1 次评审与交付说明",
                }
            )
        try:
            quote_meta = await self._doc_service.create_quote_docx(
                seller="CoDeskTeam",
                buyer="目标客户",
                currency="CNY",
                items=quote_items,
                note="内容创作与交付建议报价（可按范围调整）。",
            )
            artifact = await self._artifact_from_meta(
                db=db,
                user=user,
                meta=quote_meta,
                kind="content",
                detail="内容交付报价单（DOCX）",
                step="quotation",
                project_id=project_id,
            )
            if artifact:
                artifacts.append(artifact)
        except Exception as e:
            warnings.append(f"报价单生成失败：{e}")
        report = "\n".join(
            [
                "# 内容创作与写作 Pipeline 报告",
                "",
                f"- 主题：{topic}",
                f"- 受众：{audience}",
                f"- 模式：{mode}",
                f"- 语言：{', '.join(languages)}",
                f"- 版本数：{variants}",
                f"- 品牌词：{'、'.join(brand_terms) if brand_terms else '未设置'}",
                f"- 输入原文：{'有' if source_text else '无'}",
                f"- 输出产物数：{len(artifacts)}",
                "",
                "## 已执行",
                "- 多渠道文案包生成",
                "- 多语与多版本改写",
                "- 审阅建议与版本摘要生成",
                "- 方案文档生成",
                "- 提案 PPT 生成",
                "- 报价单文档生成",
            ]
            + (["", "## 告警", *[f"- {w}" for w in warnings]] if warnings else [])
        )
        artifacts.insert(
            0,
            await self._new_text_artifact(
                db=db,
                user=user,
                filename_hint=f"content-pipeline-report-{uuid4().hex[:8]}.md",
                text=report,
                kind="content",
                detail="内容 pipeline 执行报告",
                step="summary",
                project_id=project_id,
            ),
        )
        return report, artifacts, warnings
    async def _run_office_pipeline(
        self,
        *,
        payload: dict[str, Any],
        mode: str,
        user,
        db,
        project_id: int | None,
    ) -> tuple[str, list[PipelineArtifact], list[str]]:
        project_name = _trim(payload.get("project_name")) or "AI+通用工具集推进项目"
        tasks_raw = payload.get("tasks") if isinstance(payload.get("tasks"), list) else []
        schedule_raw = payload.get("schedule_items") if isinstance(payload.get("schedule_items"), list) else []
        reminder_hours = [max(1, _safe_int(v, 0)) for v in (payload.get("reminder_hours") or [24, 2]) if _safe_int(v, 0) > 0]
        meeting_notes = _trim(payload.get("meeting_notes"))
        conversion_targets = [str(v).strip().lower() for v in (payload.get("file_conversion_targets") or []) if str(v).strip()]
        conversion_file_ids = [str(v) for v in (payload.get("input_file_ids") or []) if str(v).strip()]
        tasks: list[dict[str, str]] = []
        for idx, item in enumerate(tasks_raw, 1):
            if not isinstance(item, dict):
                continue
            title = _trim(item.get("title")) or f"任务 {idx}"
            tasks.append(
                {
                    "title": title,
                    "owner": _trim(item.get("owner")) or "待指派",
                    "due_date": _trim(item.get("due_date")) or "TBD",
                    "priority": _trim(item.get("priority")) or "medium",
                    "status": _trim(item.get("status")) or "todo",
                }
            )
        if not tasks:
            tasks = [
                {"title": "确认需求边界", "owner": "PM", "due_date": "本周", "priority": "high", "status": "todo"},
                {"title": "搭建云边执行链路", "owner": "工程", "due_date": "下周", "priority": "high", "status": "todo"},
                {"title": "验收交付模板", "owner": "交付", "due_date": "下周", "priority": "medium", "status": "todo"},
            ]
        schedule_items: list[dict[str, Any]] = []
        for idx, item in enumerate(schedule_raw, 1):
            if not isinstance(item, dict):
                continue
            start = _parse_datetime(item.get("start"))
            end = _parse_datetime(item.get("end"))
            if start and end and end <= start:
                end = start + timedelta(hours=1)
            schedule_items.append(
                {
                    "title": _trim(item.get("title")) or f"日程 {idx}",
                    "owner": _trim(item.get("owner")) or "待定",
                    "start": start,
                    "end": end or (start + timedelta(hours=1) if start else None),
                    "priority": _trim(item.get("priority")) or "medium",
                    "location": _trim(item.get("location")) or "",
                }
            )
        if not schedule_items:
            now = datetime.utcnow().replace(minute=0, second=0, microsecond=0)
            schedule_items = [
                {"title": "需求评审", "owner": "PM", "start": now + timedelta(hours=2), "end": now + timedelta(hours=3), "priority": "high", "location": "会议室 A"},
                {"title": "技术联调", "owner": "研发", "start": now + timedelta(hours=3), "end": now + timedelta(hours=5), "priority": "high", "location": "线上"},
            ]
        warnings: list[str] = []
        artifacts: list[PipelineArtifact] = []
        csv_lines = [["title", "owner", "due_date", "priority", "status"]]
        for task in tasks:
            csv_lines.append([task["title"], task["owner"], task["due_date"], task["priority"], task["status"]])
        csv_buf: list[str] = []
        for row in csv_lines:
            csv_buf.append(",".join(json.dumps(col, ensure_ascii=False) if ("," in col or "\"" in col) else col for col in row))
        artifacts.append(
            await self._new_text_artifact(
                db=db,
                user=user,
                filename_hint=f"office-task-board-{uuid4().hex[:8]}.csv",
                text="\n".join(csv_buf),
                kind="office",
                detail="任务协作看板（CSV）",
                step="task_board",
                project_id=project_id,
            )
        )
        schedule_lines = [
            "BEGIN:VCALENDAR",
            "VERSION:2.0",
            "PRODID:-//JetLinks AI//AI Toolkit//CN",
            "CALSCALE:GREGORIAN",
        ]
        conflicts: list[str] = []
        sorted_items = sorted(schedule_items, key=lambda item: item["start"] or datetime.max)
        for idx, item in enumerate(sorted_items):
            start = item["start"]
            end = item["end"]
            uid = f"{uuid4().hex}@jetlinks-ai"
            if start and end:
                schedule_lines += [
                    "BEGIN:VEVENT",
                    f"UID:{uid}",
                    f"DTSTAMP:{datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')}",
                    f"DTSTART:{start.strftime('%Y%m%dT%H%M%S')}",
                    f"DTEND:{end.strftime('%Y%m%dT%H%M%S')}",
                    f"SUMMARY:{_escape_ics(item['title'])}",
                    f"DESCRIPTION:{_escape_ics('负责人：' + item['owner'])}",
                    f"LOCATION:{_escape_ics(item['location'])}",
                    "END:VEVENT",
                ]
            if idx > 0:
                prev = sorted_items[idx - 1]
                prev_end = prev.get("end")
                if start and prev_end and start < prev_end:
                    conflicts.append(f"{prev['title']} 与 {item['title']} 存在时间冲突（{_format_dt(prev_end)} 前重叠）。")
        schedule_lines.append("END:VCALENDAR")
        artifacts.append(
            await self._new_text_artifact(
                db=db,
                user=user,
                filename_hint=f"office-schedule-{uuid4().hex[:8]}.ics",
                text="\n".join(schedule_lines),
                kind="office",
                detail="日程安排（ICS）",
                step="schedule",
                project_id=project_id,
            )
        )
        conflict_report = "\n".join(
            [
                f"# {project_name} 日程冲突报告",
                "",
                "## 日程列表",
                *[
                    f"- {item['title']}｜负责人：{item['owner']}｜开始：{_format_dt(item['start'])}｜结束：{_format_dt(item['end'])}｜地点：{item['location'] or '未填写'}"
                    for item in sorted_items
                ],
                "",
                "## 冲突检测",
                *([f"- {item}" for item in conflicts] if conflicts else ["- 未检测到时间冲突。"]),
                "",
                "## 提醒策略",
                *[
                    f"- 在开始前 {hours} 小时提醒：{', '.join(item['title'] for item in sorted_items if item.get('start')) or '暂无可提醒事件'}"
                    for hours in reminder_hours
                ],
            ]
        )
        artifacts.append(
            await self._new_text_artifact(
                db=db,
                user=user,
                filename_hint=f"office-schedule-report-{uuid4().hex[:8]}.md",
                text=conflict_report,
                kind="office",
                detail="日程冲突与提醒报告",
                step="schedule_report",
                project_id=project_id,
            )
        )
        office_md = "\n".join(
            [
                f"# {project_name}",
                "",
                f"- 部署模式：{mode}",
                "- 输出：日程管理、任务协作、文件交付链路",
                "",
                "## 任务清单",
                *[
                    f"- {task['title']}｜负责人：{task['owner']}｜到期：{task['due_date']}｜优先级：{task['priority']}｜状态：{task['status']}"
                    for task in tasks
                ],
                "",
                "## 协作建议",
                "- 每日站会更新状态，阻塞项在 24h 内升级。",
                "- 云端用于汇总分析，边缘端负责现场执行与同步。",
            ]
        )
        artifacts.append(
            await self._new_text_artifact(
                db=db,
                user=user,
                filename_hint=f"office-playbook-{uuid4().hex[:8]}.md",
                text=office_md,
                kind="office",
                detail="办公效率执行手册",
                step="playbook",
                project_id=project_id,
            )
        )
        if meeting_notes:
            fallback_summary = "\n".join(
                [
                    f"# {project_name} 会议纪要",
                    "",
                    "## 原始记录",
                    meeting_notes,
                    "",
                    "## 摘要",
                    "已整理会议关键点，请补充负责人、截止时间与风险项。",
                ]
            )
            meeting_summary = await self._generate_text(
                system_prompt="你是一名会议纪要助手，请把输入整理为精简纪要，包含结论、行动项、风险。",
                user_prompt=meeting_notes,
                fallback=fallback_summary,
            )
            artifacts.append(
                await self._new_text_artifact(
                    db=db,
                    user=user,
                    filename_hint=f"office-meeting-summary-{uuid4().hex[:8]}.md",
                    text=meeting_summary,
                    kind="office",
                    detail="会议摘要",
                    step="meeting_summary",
                    project_id=project_id,
                )
            )
        if conversion_file_ids and conversion_targets:
            conversion_inputs, conversion_input_warnings = await self._resolve_input_files(file_ids=conversion_file_ids, user=user, db=db)
            warnings.extend(conversion_input_warnings)
            manifest_lines = ["# 文件转换结果", ""]
            for item in conversion_inputs:
                source_path: Path = item["path"]
                for fmt in conversion_targets:
                    payload_bytes, suffix, err = self._convert_file_for_office(path=source_path, target_format=fmt)
                    if payload_bytes is None or not suffix:
                        manifest_lines.append(f"- {item['filename']} -> {fmt}: 失败（{err}）")
                        warnings.append(f"文件转换失败：{item['filename']} -> {fmt} ({err})")
                        continue
                    output_name = f"{_safe_stem(source_path.stem)}-{uuid4().hex[:8]}{suffix}"
                    artifact = await self._new_artifact(
                        db=db,
                        user=user,
                        filename_hint=output_name,
                        suffix=suffix,
                        content=payload_bytes,
                        kind="office",
                        detail=f"文件转换 {fmt.upper()}",
                        step="file_convert",
                        project_id=project_id,
                    )
                    artifacts.append(artifact)
                    manifest_lines.append(f"- {item['filename']} -> {artifact.filename}: 成功")
            artifacts.append(
                await self._new_text_artifact(
                    db=db,
                    user=user,
                    filename_hint=f"office-conversion-manifest-{uuid4().hex[:8]}.md",
                    text="\n".join(manifest_lines),
                    kind="office",
                    detail="文件转换清单",
                    step="file_convert_manifest",
                    project_id=project_id,
                )
            )
        quote_items = [
            {"name": task["title"], "quantity": 1, "unit_price": float(2000 + idx * 800), "unit": "项", "note": "执行与追踪"}
            for idx, task in enumerate(tasks, 1)
        ]
        try:
            quote_meta = await self._doc_service.create_quote_xlsx(
                seller="CoDeskTeam",
                buyer=project_name,
                currency="CNY",
                items=quote_items,
                note="办公效率与协作能力建设估算清单。",
            )
            artifact = await self._artifact_from_meta(
                db=db,
                user=user,
                meta=quote_meta,
                kind="office",
                detail="办公效率报价单（XLSX）",
                step="quotation",
                project_id=project_id,
            )
            if artifact:
                artifacts.append(artifact)
        except Exception as e:
            warnings.append(f"XLSX 报价生成失败：{e}")
        try:
            proto_meta = await self._prototype_service.generate(
                project_name=f"{project_name} 协作原型",
                pages=[
                    {"title": "日程管理", "description": "周计划与冲突检查"},
                    {"title": "任务协作", "description": "任务拆解与进度跟踪"},
                    {"title": "文件中心", "description": "格式转换与交付记录"},
                ],
            )
            artifact = await self._artifact_from_meta(
                db=db,
                user=user,
                meta=proto_meta,
                kind="office",
                detail="协作原型 ZIP",
                step="prototype",
                project_id=project_id,
            )
            if artifact:
                artifacts.append(artifact)
        except Exception as e:
            warnings.append(f"协作原型生成失败：{e}")
        report = "\n".join(
            [
                "# 办公与效率提升 Pipeline 报告",
                "",
                f"- 项目：{project_name}",
                f"- 模式：{mode}",
                f"- 任务数：{len(tasks)}",
                f"- 日程数：{len(schedule_items)}",
                f"- 提醒策略：{', '.join(f'{v}h' for v in reminder_hours)}",
                f"- 会议摘要：{'已生成' if meeting_notes else '未输入'}",
                f"- 文件转换目标：{', '.join(conversion_targets) if conversion_targets else '未启用'}",
                f"- 输出产物数：{len(artifacts)}",
                "",
                "## 已执行",
                "- 日程导出与冲突检测",
                "- 任务协作看板导出",
                "- 办公执行手册生成",
                "- 会议摘要与提醒策略生成",
                "- 文件转换清单输出",
                "- 报价与原型交付文件生成",
            ]
            + (["", "## 告警", *[f"- {w}" for w in warnings]] if warnings else [])
        )
        artifacts.insert(
            0,
            await self._new_text_artifact(
                db=db,
                user=user,
                filename_hint=f"office-pipeline-report-{uuid4().hex[:8]}.md",
                text=report,
                kind="office",
                detail="办公 pipeline 执行报告",
                step="summary",
                project_id=project_id,
            )
        )
        return report, artifacts, warnings
    async def _run_full_pipeline(
        self,
        *,
        payload: dict[str, Any],
        mode: str,
        user,
        db,
        project_id: int | None,
    ) -> tuple[str, list[PipelineArtifact], list[str]]:
        selected_raw = payload.get("pipelines")
        selected = [str(v).strip().lower() for v in (selected_raw or []) if str(v).strip()]
        if not selected:
            selected = ["vision", "media", "content", "office"]
        sections: list[tuple[str, str, dict[str, Any]]] = []
        for key in selected:
            if key == "vision":
                sections.append(("视觉", "vision", payload.get("vision_payload") if isinstance(payload.get("vision_payload"), dict) else {}))
            elif key == "media":
                sections.append(("音视频", "media", payload.get("media_payload") if isinstance(payload.get("media_payload"), dict) else {}))
            elif key == "content":
                sections.append(("内容", "content", payload.get("content_payload") if isinstance(payload.get("content_payload"), dict) else {}))
            elif key == "office":
                sections.append(("办公", "office", payload.get("office_payload") if isinstance(payload.get("office_payload"), dict) else {}))
        all_artifacts: list[PipelineArtifact] = []
        all_warnings: list[str] = []
        section_summaries: list[str] = []
        for label, key, section_payload in sections:
            merged = {
                **section_payload,
                "deployment_mode": mode,
            }
            if key in {"vision", "media"} and not merged.get("input_file_ids"):
                merged["input_file_ids"] = payload.get("input_file_ids") if isinstance(payload.get("input_file_ids"), list) else []
            if key == "vision":
                summary, artifacts, warnings = await self._run_vision_pipeline(
                    payload=merged, mode=mode, user=user, db=db, project_id=project_id
                )
            elif key == "media":
                summary, artifacts, warnings = await self._run_media_pipeline(
                    payload=merged, mode=mode, user=user, db=db, project_id=project_id
                )
            elif key == "content":
                summary, artifacts, warnings = await self._run_content_pipeline(
                    payload=merged, mode=mode, user=user, db=db, project_id=project_id
                )
            else:
                summary, artifacts, warnings = await self._run_office_pipeline(
                    payload=merged, mode=mode, user=user, db=db, project_id=project_id
                )
            section_summaries.append(f"## {label}\n\n{summary}")
            all_artifacts.extend(artifacts)
            all_warnings.extend(warnings)
        orchestration = "\n".join(
            [
                "# AI+通用工具集全链路 Pipeline 报告",
                "",
                f"- 部署模式：{mode}",
                f"- 覆盖流程：{', '.join(label for label, _, _ in sections)}",
                f"- 总产物数：{len(all_artifacts)}",
                "",
                "## 汇总建议",
                "- 云端承担高复杂生成与批处理。",
                "- 边缘端承担低时延处理与本地闭环。",
                "- 使用统一版本命名和回滚策略保证可治理执行。",
                "",
                *section_summaries,
            ]
        )
        if all_warnings:
            orchestration += "\n\n## 全局告警\n" + "\n".join(f"- {w}" for w in all_warnings)
        all_artifacts.insert(
            0,
            await self._new_text_artifact(
                db=db,
                user=user,
                filename_hint=f"toolkit-full-pipeline-report-{uuid4().hex[:8]}.md",
                text=orchestration,
                kind="full",
                detail="全链路 pipeline 汇总报告",
                step="summary",
                project_id=project_id,
            ),
        )
        return orchestration, all_artifacts, all_warnings
